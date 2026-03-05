from fastapi import FastAPI, APIRouter, HTTPException, Response, Request, UploadFile, File, Form
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import httpx
import base64
import json
import asyncio

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app
app = FastAPI(title="Criminal Appeal AI - Case Management")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============ ROOT HEALTH CHECK (for Kubernetes) ============
@app.get("/health")
async def root_health_check():
    """Root-level health check for Kubernetes deployment"""
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

# ============ MODELS ============

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Case(BaseModel):
    model_config = ConfigDict(extra="ignore")
    case_id: str = Field(default_factory=lambda: f"case_{uuid.uuid4().hex[:12]}")
    user_id: str
    title: str
    defendant_name: str
    case_number: Optional[str] = None
    court: Optional[str] = None
    judge: Optional[str] = None
    status: str = "active"
    summary: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CaseCreate(BaseModel):
    title: str
    defendant_name: str
    case_number: Optional[str] = None
    court: Optional[str] = None
    judge: Optional[str] = None
    summary: Optional[str] = None

class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    document_id: str = Field(default_factory=lambda: f"doc_{uuid.uuid4().hex[:12]}")
    case_id: str
    user_id: str
    filename: str
    file_type: str
    category: str  # brief, case_note, public_advertising, evidence, court_document, other
    description: Optional[str] = None
    content_text: Optional[str] = None  # Extracted text for analysis
    file_data: Optional[str] = None  # Base64 encoded file data
    event_date: Optional[datetime] = None
    uploaded_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TimelineEvent(BaseModel):
    model_config = ConfigDict(extra="ignore")
    event_id: str = Field(default_factory=lambda: f"evt_{uuid.uuid4().hex[:12]}")
    case_id: str
    user_id: str
    title: str
    description: str
    event_date: datetime
    event_type: str  # See EVENT_CATEGORIES below
    event_category: str = "general"  # pre_trial, trial, evidence, post_conviction, investigation
    
    # Enhanced fields
    linked_documents: List[str] = []  # Document IDs
    participants: List[dict] = []  # [{name: str, role: str}]
    significance: str = "normal"  # critical, important, normal, minor
    source_citation: str = ""
    perspective: str = "neutral"  # prosecution, defence, neutral
    is_contested: bool = False
    contested_details: str = ""
    related_grounds: List[str] = []  # Ground IDs
    inconsistency_notes: str = ""
    
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TimelineEventCreate(BaseModel):
    title: str
    description: str
    event_date: datetime
    event_type: str
    event_category: str = "general"
    linked_documents: List[str] = []
    participants: List[dict] = []
    significance: str = "normal"
    source_citation: str = ""
    perspective: str = "neutral"
    is_contested: bool = False
    contested_details: str = ""
    related_grounds: List[str] = []
    inconsistency_notes: str = ""

class GroundOfMerit(BaseModel):
    model_config = ConfigDict(extra="ignore")
    ground_id: str = Field(default_factory=lambda: f"gnd_{uuid.uuid4().hex[:12]}")
    case_id: str
    user_id: str
    title: str
    ground_type: str  # procedural_error, fresh_evidence, miscarriage_of_justice, sentencing_error, judicial_error, ineffective_counsel, other
    description: str
    strength: str = "moderate"  # strong, moderate, weak
    status: str = "identified"  # identified, investigating, confirmed, rejected
    supporting_evidence: List[str] = []
    law_sections: List[dict] = []  # {jurisdiction, section, title, relevance, full_text}
    similar_cases: List[dict] = []  # {case_name, citation, relevance, outcome}
    analysis: Optional[str] = None
    deep_analysis: Optional[dict] = None  # Detailed analysis content
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class GroundOfMeritCreate(BaseModel):
    title: str
    ground_type: str = "other"
    description: str
    strength: str = "moderate"
    supporting_evidence: List[str] = []

class GroundOfMeritUpdate(BaseModel):
    title: Optional[str] = None
    ground_type: Optional[str] = None
    description: Optional[str] = None
    strength: Optional[str] = None
    status: Optional[str] = None
    supporting_evidence: Optional[List[str]] = None

class ReportRequest(BaseModel):
    report_type: str = "quick_summary"

class Report(BaseModel):
    model_config = ConfigDict(extra="ignore")
    report_id: str = Field(default_factory=lambda: f"rpt_{uuid.uuid4().hex[:12]}")
    case_id: str
    user_id: str
    report_type: str  # quick_summary, full_detailed, extensive_log
    title: str
    content: dict  # JSON structure with report sections
    grounds_of_merit: List[dict] = []
    generated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Note(BaseModel):
    model_config = ConfigDict(extra="ignore")
    note_id: str = Field(default_factory=lambda: f"note_{uuid.uuid4().hex[:12]}")
    case_id: str
    user_id: str
    author_name: str
    author_email: str
    category: str = "general"  # general, legal_opinion, evidence_note, strategy, question, action_item
    title: str
    content: str
    is_pinned: bool = False
    document_id: Optional[str] = None  # Link to specific document
    report_id: Optional[str] = None  # Link to specific report
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class NoteCreate(BaseModel):
    category: str = "general"
    title: str
    content: str
    is_pinned: bool = False
    document_id: Optional[str] = None
    report_id: Optional[str] = None

class NoteUpdate(BaseModel):
    category: Optional[str] = None
    title: Optional[str] = None
    content: Optional[str] = None
    is_pinned: Optional[bool] = None

# ============ DEADLINE MODELS ============

class Deadline(BaseModel):
    model_config = ConfigDict(extra="ignore")
    deadline_id: str = Field(default_factory=lambda: f"dl_{uuid.uuid4().hex[:12]}")
    case_id: str
    user_id: str
    title: str
    description: str = ""
    deadline_type: str  # appeal_lodgement, leave_application, document_filing, hearing, other
    due_date: datetime
    reminder_days: List[int] = [7, 3, 1]  # Days before to remind
    is_completed: bool = False
    completed_at: Optional[datetime] = None
    priority: str = "high"  # critical, high, medium, low
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class DeadlineCreate(BaseModel):
    title: str
    description: str = ""
    deadline_type: str = "other"
    due_date: datetime
    reminder_days: List[int] = [7, 3, 1]
    priority: str = "high"

# ============ CHECKLIST MODELS ============

class ChecklistItem(BaseModel):
    model_config = ConfigDict(extra="ignore")
    item_id: str = Field(default_factory=lambda: f"chk_{uuid.uuid4().hex[:12]}")
    case_id: str
    user_id: str
    phase: str  # preparation, grounds_identification, investigation, documentation, lodgement, hearing
    title: str
    description: str = ""
    is_completed: bool = False
    completed_at: Optional[datetime] = None
    order: int = 0
    is_custom: bool = False  # True if user-added, False if default
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Default checklist items for new cases
DEFAULT_CHECKLIST = [
    # Phase 1: Preparation
    {"phase": "preparation", "title": "Gather all case documents", "description": "Collect court transcripts, evidence briefs, witness statements, and any other relevant documents", "order": 1},
    {"phase": "preparation", "title": "Upload documents to system", "description": "Upload all gathered documents and extract text for AI analysis", "order": 2},
    {"phase": "preparation", "title": "Build case timeline", "description": "Create chronological timeline of events from arrest to conviction", "order": 3},
    {"phase": "preparation", "title": "Review trial transcript", "description": "Read through the full trial transcript noting any issues", "order": 4},
    # Phase 2: Grounds Identification
    {"phase": "grounds_identification", "title": "Run AI grounds identification", "description": "Use AI to identify potential grounds of merit from your documents", "order": 5},
    {"phase": "grounds_identification", "title": "Review identified grounds", "description": "Assess each ground identified by AI for viability", "order": 6},
    {"phase": "grounds_identification", "title": "Check witness statement inconsistencies", "description": "Use contradiction finder to identify issues in witness testimony", "order": 7},
    {"phase": "grounds_identification", "title": "Identify any fresh evidence", "description": "Document any new evidence not available at trial", "order": 8},
    # Phase 3: Investigation
    {"phase": "investigation", "title": "Deep investigate each ground", "description": "Use AI investigation feature on each viable ground", "order": 9},
    {"phase": "investigation", "title": "Research relevant case law", "description": "Find similar appeal cases and their outcomes", "order": 10},
    {"phase": "investigation", "title": "Identify relevant law sections", "description": "Document NSW and Federal law sections that apply", "order": 11},
    {"phase": "investigation", "title": "Assess case strength", "description": "Review case strength meter and prioritize strongest grounds", "order": 12},
    # Phase 4: Documentation
    {"phase": "documentation", "title": "Generate detailed report", "description": "Create Full Detailed report with all grounds and analysis", "order": 13},
    {"phase": "documentation", "title": "Prepare Notice of Appeal", "description": "Draft Notice of Appeal using template", "order": 14},
    {"phase": "documentation", "title": "Prepare supporting affidavit", "description": "Draft affidavit if required for fresh evidence", "order": 15},
    {"phase": "documentation", "title": "Review all documents", "description": "Final review of all appeal documents for accuracy", "order": 16},
    # Phase 5: Lodgement
    {"phase": "lodgement", "title": "Consult with legal professional", "description": "Have documents reviewed by solicitor or barrister", "order": 17},
    {"phase": "lodgement", "title": "File Notice of Appeal", "description": "Lodge appeal within 28 days of conviction/sentence", "order": 18},
    {"phase": "lodgement", "title": "File Leave application (if required)", "description": "Lodge application for leave to appeal if needed", "order": 19},
    {"phase": "lodgement", "title": "Confirm lodgement receipt", "description": "Obtain confirmation that appeal has been filed", "order": 20},
    # Phase 6: Hearing Preparation
    {"phase": "hearing", "title": "Prepare for hearing", "description": "Review all materials before appeal hearing", "order": 21},
    {"phase": "hearing", "title": "Generate Barrister View report", "description": "Create professional presentation for court", "order": 22},
]

# ============ PAYMENT MODELS ============

class Payment(BaseModel):
    model_config = ConfigDict(extra="ignore")
    payment_id: str = Field(default_factory=lambda: f"pay_{uuid.uuid4().hex[:12]}")
    user_id: str
    case_id: str
    feature_type: str  # "grounds_of_merit", "full_report", "extensive_report"
    amount: float
    currency: str = "AUD"
    paypal_order_id: Optional[str] = None
    status: str = "pending"  # pending, completed, failed, refunded
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    completed_at: Optional[datetime] = None

# Payment pricing configuration
FEATURE_PRICES = {
    "grounds_of_merit": {"price": 50.00, "name": "Unlock Grounds of Merit Details"},
    "full_report": {"price": 29.99, "name": "Full Detailed Report"},
    "extensive_report": {"price": 50.00, "name": "Extensive Log Report"}
}

# PayPal Configuration
import paypalrestsdk
PAYPAL_CLIENT_ID = os.environ.get('PAYPAL_CLIENT_ID', '')
PAYPAL_CLIENT_SECRET = os.environ.get('PAYPAL_CLIENT_SECRET', '')
PAYPAL_MODE = os.environ.get('PAYPAL_MODE', 'sandbox')  # sandbox or live
PAYPAL_CONFIGURED = False

if PAYPAL_CLIENT_ID and PAYPAL_CLIENT_SECRET:
    try:
        paypalrestsdk.configure({
            "mode": PAYPAL_MODE,
            "client_id": PAYPAL_CLIENT_ID,
            "client_secret": PAYPAL_CLIENT_SECRET
        })
        PAYPAL_CONFIGURED = True
        logger.info(f"PayPal configured in {PAYPAL_MODE} mode")
    except Exception as e:
        logger.warning(f"PayPal configuration failed: {e}")
        PAYPAL_CONFIGURED = False
else:
    logger.info("PayPal credentials not configured - payment features disabled")

# ============ AUTH HELPERS ============

async def get_current_user(request: Request) -> User:
    """Get current user from session token (cookie or header)"""
    session_token = request.cookies.get("session_token")
    if not session_token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            session_token = auth_header.split(" ")[1]
    
    if not session_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    session_doc = await db.user_sessions.find_one({"session_token": session_token}, {"_id": 0})
    if not session_doc:
        raise HTTPException(status_code=401, detail="Invalid session")
    
    # Check expiry
    expires_at = session_doc.get("expires_at")
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Session expired")
    
    user_doc = await db.users.find_one({"user_id": session_doc["user_id"]}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=401, detail="User not found")
    
    return User(**user_doc)

# ============ AUTH ENDPOINTS ============

@api_router.post("/auth/session")
async def create_session(request: Request, response: Response):
    """Exchange session_id for session_token"""
    body = await request.json()
    session_id = body.get("session_id")
    
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id required")
    
    # Call Emergent Auth to get user data
    async with httpx.AsyncClient() as client:
        auth_response = await client.get(
            "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
            headers={"X-Session-ID": session_id}
        )
    
    if auth_response.status_code != 200:
        raise HTTPException(status_code=401, detail="Invalid session_id")
    
    user_data = auth_response.json()
    
    # Create or update user
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    existing_user = await db.users.find_one({"email": user_data["email"]}, {"_id": 0})
    
    if existing_user:
        user_id = existing_user["user_id"]
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {
                "name": user_data["name"],
                "picture": user_data.get("picture"),
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
    else:
        await db.users.insert_one({
            "user_id": user_id,
            "email": user_data["email"],
            "name": user_data["name"],
            "picture": user_data.get("picture"),
            "created_at": datetime.now(timezone.utc).isoformat()
        })
    
    # Create session
    session_token = user_data.get("session_token", f"sess_{uuid.uuid4().hex}")
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Set cookie
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7 * 24 * 60 * 60
    )
    
    user_doc = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    return user_doc

@api_router.get("/auth/me")
async def get_me(request: Request):
    """Get current user info"""
    user = await get_current_user(request)
    user_dict = user.model_dump()
    # Check if user has accepted terms
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    user_dict["terms_accepted"] = user_doc.get("terms_accepted", False)
    user_dict["terms_accepted_at"] = user_doc.get("terms_accepted_at")
    return user_dict

@api_router.post("/auth/accept-terms")
async def accept_terms(request: Request):
    """Accept terms and conditions"""
    user = await get_current_user(request)
    await db.users.update_one(
        {"user_id": user.user_id},
        {"$set": {
            "terms_accepted": True,
            "terms_accepted_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    return {"message": "Terms accepted", "terms_accepted": True}

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    """Logout user"""
    session_token = request.cookies.get("session_token")
    if session_token:
        await db.user_sessions.delete_one({"session_token": session_token})
    
    response.delete_cookie(key="session_token", path="/", samesite="none", secure=True)
    return {"message": "Logged out"}

# ============ CASE ENDPOINTS ============

@api_router.get("/cases", response_model=List[dict])
async def get_cases(request: Request):
    """Get all cases for current user"""
    user = await get_current_user(request)
    cases = await db.cases.find({"user_id": user.user_id}, {"_id": 0}).sort("updated_at", -1).to_list(100)
    
    if not cases:
        return cases
    
    # Optimize: Batch count queries using aggregation instead of N+1 pattern
    case_ids = [case["case_id"] for case in cases]
    
    # Get document counts in a single aggregation query
    doc_counts = await db.documents.aggregate([
        {"$match": {"case_id": {"$in": case_ids}}},
        {"$group": {"_id": "$case_id", "count": {"$sum": 1}}}
    ]).to_list(100)
    
    # Get event counts in a single aggregation query
    event_counts = await db.timeline_events.aggregate([
        {"$match": {"case_id": {"$in": case_ids}}},
        {"$group": {"_id": "$case_id", "count": {"$sum": 1}}}
    ]).to_list(100)
    
    # Map counts back to cases
    doc_map = {d["_id"]: d["count"] for d in doc_counts}
    event_map = {e["_id"]: e["count"] for e in event_counts}
    
    for case in cases:
        case["document_count"] = doc_map.get(case["case_id"], 0)
        case["event_count"] = event_map.get(case["case_id"], 0)
    
    return cases

@api_router.post("/cases", response_model=dict)
async def create_case(case_data: CaseCreate, request: Request):
    """Create a new case"""
    user = await get_current_user(request)
    
    case = Case(
        user_id=user.user_id,
        **case_data.model_dump()
    )
    
    case_dict = case.model_dump()
    case_dict["created_at"] = case_dict["created_at"].isoformat()
    case_dict["updated_at"] = case_dict["updated_at"].isoformat()
    
    await db.cases.insert_one(case_dict)
    
    # Return the case without MongoDB's _id field
    created_case = await db.cases.find_one({"case_id": case.case_id}, {"_id": 0})
    return created_case

@api_router.get("/cases/{case_id}", response_model=dict)
async def get_case(case_id: str, request: Request):
    """Get a specific case"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Add counts
    case["document_count"] = await db.documents.count_documents({"case_id": case_id})
    case["event_count"] = await db.timeline_events.count_documents({"case_id": case_id})
    
    return case

@api_router.put("/cases/{case_id}", response_model=dict)
async def update_case(case_id: str, case_data: CaseCreate, request: Request):
    """Update a case"""
    user = await get_current_user(request)
    
    result = await db.cases.update_one(
        {"case_id": case_id, "user_id": user.user_id},
        {"$set": {
            **case_data.model_dump(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Case not found")
    
    return await db.cases.find_one({"case_id": case_id}, {"_id": 0})

@api_router.delete("/cases/{case_id}")
async def delete_case(case_id: str, request: Request):
    """Delete a case and all related data"""
    user = await get_current_user(request)
    
    # Delete related data
    await db.documents.delete_many({"case_id": case_id, "user_id": user.user_id})
    await db.timeline_events.delete_many({"case_id": case_id, "user_id": user.user_id})
    await db.reports.delete_many({"case_id": case_id, "user_id": user.user_id})
    
    result = await db.cases.delete_one({"case_id": case_id, "user_id": user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Case not found")
    
    return {"message": "Case deleted"}

# ============ DOCUMENT ENDPOINTS ============

@api_router.get("/cases/{case_id}/documents", response_model=List[dict])
async def get_documents(case_id: str, request: Request):
    """Get all documents for a case"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0, "file_data": 0}  # Exclude file data for listing
    ).sort("uploaded_at", -1).to_list(500)
    
    return documents

@api_router.post("/cases/{case_id}/documents", response_model=dict)
async def upload_document(
    case_id: str,
    request: Request,
    file: UploadFile = File(...),
    category: str = Form(...),
    description: str = Form(None),
    event_date: str = Form(None)
):
    """Upload a document to a case"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Read file content
    file_content = await file.read()
    file_base64 = base64.b64encode(file_content).decode('utf-8')
    
    # Extract text content based on file type
    content_text = ""
    file_type = file.content_type or "application/octet-stream"
    filename_lower = file.filename.lower() if file.filename else ""
    
    try:
        if "text" in file_type or filename_lower.endswith('.txt'):
            content_text = file_content.decode('utf-8', errors='ignore')
        
        elif "pdf" in file_type or filename_lower.endswith('.pdf'):
            # Extract text from PDF
            try:
                import io
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page in pdf_reader.pages[:20]:  # Limit to first 20 pages
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
                content_text = ""
        
        elif filename_lower.endswith('.docx') or "word" in file_type:
            # Extract text from DOCX
            try:
                import io
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(file_content))
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
                content_text = ""
    except Exception as e:
        logger.warning(f"Text extraction error: {e}")
        content_text = ""
    
    # Parse event date
    parsed_event_date = None
    if event_date:
        try:
            parsed_event_date = datetime.fromisoformat(event_date.replace('Z', '+00:00')).isoformat()
        except ValueError:
            pass
    
    doc = Document(
        case_id=case_id,
        user_id=user.user_id,
        filename=file.filename,
        file_type=file_type,
        category=category,
        description=description,
        content_text=content_text,
        file_data=file_base64
    )
    
    doc_dict = doc.model_dump()
    doc_dict["uploaded_at"] = doc_dict["uploaded_at"].isoformat()
    if parsed_event_date:
        doc_dict["event_date"] = parsed_event_date
    
    await db.documents.insert_one(doc_dict)
    
    # Update case updated_at
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Return the document without MongoDB's _id field and file_data
    created_doc = await db.documents.find_one({"document_id": doc.document_id}, {"_id": 0, "file_data": 0})
    return created_doc

@api_router.get("/cases/{case_id}/documents/{document_id}", response_model=dict)
async def get_document(case_id: str, document_id: str, request: Request):
    """Get a specific document"""
    user = await get_current_user(request)
    
    doc = await db.documents.find_one(
        {"document_id": document_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return doc

@api_router.delete("/cases/{case_id}/documents/{document_id}")
async def delete_document(case_id: str, document_id: str, request: Request):
    """Delete a document"""
    user = await get_current_user(request)
    
    result = await db.documents.delete_one({
        "document_id": document_id,
        "case_id": case_id,
        "user_id": user.user_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {"message": "Document deleted"}

class DocumentSearchRequest(BaseModel):
    query: str
    case_sensitive: bool = False

class SearchMatch(BaseModel):
    document_id: str
    filename: str
    category: str
    matches: List[dict]  # List of {context, position}
    match_count: int

@api_router.post("/cases/{case_id}/documents/search", response_model=dict)
async def search_documents(case_id: str, search_request: DocumentSearchRequest, request: Request):
    """Search for text across all documents in a case"""
    import re
    
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    query = search_request.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Search query is required")
    
    if len(query) < 2:
        raise HTTPException(status_code=400, detail="Search query must be at least 2 characters")
    
    # Get all documents with content
    documents = await db.documents.find(
        {"case_id": case_id, "content_text": {"$exists": True, "$ne": ""}},
        {"_id": 0, "file_data": 0}
    ).to_list(500)
    
    results = []
    total_matches = 0
    
    # Search flags
    flags = 0 if search_request.case_sensitive else re.IGNORECASE
    
    for doc in documents:
        content = doc.get("content_text", "")
        if not content:
            continue
        
        matches = []
        
        # Find all occurrences with context
        try:
            pattern = re.compile(re.escape(query), flags)
            for match in pattern.finditer(content):
                start_pos = match.start()
                end_pos = match.end()
                
                # Get context (100 chars before and after)
                context_start = max(0, start_pos - 100)
                context_end = min(len(content), end_pos + 100)
                
                context = content[context_start:context_end]
                
                # Add ellipsis if truncated
                if context_start > 0:
                    context = "..." + context
                if context_end < len(content):
                    context = context + "..."
                
                matches.append({
                    "context": context,
                    "position": start_pos,
                    "matched_text": match.group()
                })
                
                # Limit matches per document to prevent huge responses
                if len(matches) >= 10:
                    break
        except re.error:
            continue
        
        if matches:
            total_matches += len(matches)
            results.append({
                "document_id": doc.get("document_id"),
                "filename": doc.get("filename"),
                "category": doc.get("category"),
                "matches": matches,
                "match_count": len(matches)
            })
    
    # Sort by match count (most matches first)
    results.sort(key=lambda x: x["match_count"], reverse=True)
    
    return {
        "query": query,
        "total_matches": total_matches,
        "documents_with_matches": len(results),
        "total_documents_searched": len(documents),
        "results": results
    }

# ============ OCR FUNCTIONS ============

def extract_text_with_ocr(file_content: bytes, filename: str, file_type: str) -> tuple:
    """Extract text from images and scanned PDFs using OCR"""
    import io
    import pytesseract
    from PIL import Image
    
    filename_lower = filename.lower()
    extracted_text = ""
    ocr_used = False
    
    try:
        # Handle image files directly
        if any(filename_lower.endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.webp']):
            image = Image.open(io.BytesIO(file_content))
            # Convert to RGB if necessary
            if image.mode in ('RGBA', 'P'):
                image = image.convert('RGB')
            extracted_text = pytesseract.image_to_string(image)
            ocr_used = True
            logger.info(f"OCR extracted {len(extracted_text)} chars from image {filename}")
        
        # Handle PDFs - try regular extraction first, then OCR if needed
        elif "pdf" in file_type or filename_lower.endswith('.pdf'):
            # First try regular PDF text extraction
            try:
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page in pdf_reader.pages[:30]:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                extracted_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"Regular PDF extraction failed: {e}")
            
            # If minimal text extracted, try OCR
            if len(extracted_text.strip()) < 100:
                try:
                    from pdf2image import convert_from_bytes
                    
                    # Convert PDF pages to images
                    images = convert_from_bytes(
                        file_content, 
                        dpi=200,
                        first_page=1,
                        last_page=min(20, 20)  # Limit to 20 pages for performance
                    )
                    
                    ocr_parts = []
                    for i, image in enumerate(images):
                        # Convert to RGB if necessary
                        if image.mode in ('RGBA', 'P'):
                            image = image.convert('RGB')
                        page_text = pytesseract.image_to_string(image)
                        if page_text.strip():
                            ocr_parts.append(f"--- Page {i+1} ---\n{page_text}")
                    
                    if ocr_parts:
                        extracted_text = "\n\n".join(ocr_parts)
                        ocr_used = True
                        logger.info(f"OCR extracted {len(extracted_text)} chars from scanned PDF {filename}")
                except Exception as e:
                    logger.warning(f"PDF OCR failed: {e}")
        
        # Handle DOCX
        elif filename_lower.endswith('.docx') or "word" in file_type:
            try:
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(io.BytesIO(file_content))
                text_parts = []
                for para in docx_doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                extracted_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
        
        # Handle plain text
        elif "text" in file_type or filename_lower.endswith('.txt'):
            extracted_text = file_content.decode('utf-8', errors='ignore')
    
    except Exception as e:
        logger.error(f"OCR/Text extraction error for {filename}: {e}")
    
    return extracted_text, ocr_used

@api_router.post("/cases/{case_id}/documents/{document_id}/ocr", response_model=dict)
async def ocr_document(case_id: str, document_id: str, request: Request):
    """Extract text from a document using OCR (for scanned documents and images)"""
    user = await get_current_user(request)
    
    doc = await db.documents.find_one(
        {"document_id": document_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not doc.get('file_data'):
        raise HTTPException(status_code=400, detail="No file data available")
    
    # Decode file content
    file_content = base64.b64decode(doc['file_data'])
    filename = doc.get('filename', '')
    file_type = doc.get('file_type', '')
    
    # Extract text using OCR
    content_text, ocr_used = extract_text_with_ocr(file_content, filename, file_type)
    
    if not content_text.strip():
        return {
            "document_id": document_id,
            "filename": filename,
            "success": False,
            "ocr_used": ocr_used,
            "message": "No text could be extracted from this document",
            "content_length": 0
        }
    
    # Update document with extracted text
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {
            "content_text": content_text,
            "ocr_extracted": ocr_used,
            "text_extracted_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "document_id": document_id,
        "filename": filename,
        "success": True,
        "ocr_used": ocr_used,
        "content_length": len(content_text),
        "content_preview": content_text[:500] + "..." if len(content_text) > 500 else content_text
    }

@api_router.post("/cases/{case_id}/ocr-all", response_model=dict)
async def ocr_all_documents(case_id: str, request: Request):
    """Run OCR on all documents in a case that don't have extracted text"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get documents without text or with minimal text
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(500)
    
    results = []
    successful_ocr = 0
    
    for doc in documents:
        existing_text = doc.get('content_text', '')
        
        # Skip if already has substantial text
        if len(existing_text.strip()) > 100:
            results.append({
                "document_id": doc['document_id'],
                "filename": doc.get('filename'),
                "status": "skipped",
                "reason": "Already has extracted text"
            })
            continue
        
        if not doc.get('file_data'):
            results.append({
                "document_id": doc['document_id'],
                "filename": doc.get('filename'),
                "status": "skipped",
                "reason": "No file data"
            })
            continue
        
        # Try OCR extraction
        file_content = base64.b64decode(doc['file_data'])
        content_text, ocr_used = extract_text_with_ocr(
            file_content, 
            doc.get('filename', ''), 
            doc.get('file_type', '')
        )
        
        if content_text.strip():
            await db.documents.update_one(
                {"document_id": doc['document_id']},
                {"$set": {
                    "content_text": content_text,
                    "ocr_extracted": ocr_used,
                    "text_extracted_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            successful_ocr += 1
            results.append({
                "document_id": doc['document_id'],
                "filename": doc.get('filename'),
                "status": "success",
                "ocr_used": ocr_used,
                "content_length": len(content_text)
            })
        else:
            results.append({
                "document_id": doc['document_id'],
                "filename": doc.get('filename'),
                "status": "failed",
                "reason": "No text could be extracted"
            })
    
    return {
        "total_documents": len(documents),
        "successful_extractions": successful_ocr,
        "results": results
    }

@api_router.post("/cases/{case_id}/documents/{document_id}/extract-text", response_model=dict)
async def extract_document_text(case_id: str, document_id: str, request: Request):
    """Re-extract text from a document (useful if initial extraction failed)"""
    user = await get_current_user(request)
    
    doc = await db.documents.find_one(
        {"document_id": document_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not doc.get('file_data'):
        raise HTTPException(status_code=400, detail="No file data available")
    
    # Decode file content
    file_content = base64.b64decode(doc['file_data'])
    filename_lower = doc.get('filename', '').lower()
    file_type = doc.get('file_type', '')
    
    content_text = ""
    
    try:
        if "text" in file_type or filename_lower.endswith('.txt'):
            content_text = file_content.decode('utf-8', errors='ignore')
        
        elif "pdf" in file_type or filename_lower.endswith('.pdf'):
            try:
                import io
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page in pdf_reader.pages[:30]:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
        
        elif filename_lower.endswith('.docx') or "word" in file_type:
            try:
                import io
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(io.BytesIO(file_content))
                text_parts = []
                for para in docx_doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {str(e)}")
    
    # Update document with extracted text
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {
            "content_text": content_text,
            "text_extracted_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "document_id": document_id,
        "filename": doc.get('filename'),
        "content_length": len(content_text),
        "content_preview": content_text[:500] + "..." if len(content_text) > 500 else content_text,
        "success": bool(content_text)
    }

@api_router.post("/cases/{case_id}/extract-all-text", response_model=dict)
async def extract_all_documents_text(case_id: str, request: Request):
    """Extract text from all documents in a case"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(500)
    
    results = []
    for doc in documents:
        if not doc.get('file_data'):
            results.append({"document_id": doc['document_id'], "success": False, "error": "No file data"})
            continue
        
        file_content = base64.b64decode(doc['file_data'])
        filename_lower = doc.get('filename', '').lower()
        file_type = doc.get('file_type', '')
        
        content_text = ""
        error = None
        
        try:
            if "text" in file_type or filename_lower.endswith('.txt'):
                content_text = file_content.decode('utf-8', errors='ignore')
            
            elif "pdf" in file_type or filename_lower.endswith('.pdf'):
                try:
                    import io
                    from PyPDF2 import PdfReader
                    pdf_reader = PdfReader(io.BytesIO(file_content))
                    text_parts = []
                    for page in pdf_reader.pages[:30]:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    content_text = "\n".join(text_parts)
                except Exception as e:
                    error = str(e)
            
            elif filename_lower.endswith('.docx') or "word" in file_type:
                try:
                    import io
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(io.BytesIO(file_content))
                    text_parts = []
                    for para in docx_doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    content_text = "\n".join(text_parts)
                except Exception as e:
                    error = str(e)
        except Exception as e:
            error = str(e)
        
        # Update document
        if content_text:
            await db.documents.update_one(
                {"document_id": doc['document_id']},
                {"$set": {
                    "content_text": content_text,
                    "text_extracted_at": datetime.now(timezone.utc).isoformat()
                }}
            )
        
        results.append({
            "document_id": doc['document_id'],
            "filename": doc.get('filename'),
            "success": bool(content_text),
            "content_length": len(content_text) if content_text else 0,
            "error": error
        })
    
    successful = sum(1 for r in results if r['success'])
    return {
        "total_documents": len(documents),
        "successful_extractions": successful,
        "results": results
    }

# ============ TIMELINE ENDPOINTS ============

@api_router.get("/cases/{case_id}/timeline", response_model=List[dict])
async def get_timeline(case_id: str, request: Request):
    """Get timeline events for a case"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    events = await db.timeline_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("event_date", 1).to_list(500)
    
    return events

@api_router.post("/cases/{case_id}/timeline", response_model=dict)
async def create_timeline_event(case_id: str, event_data: TimelineEventCreate, request: Request):
    """Create a timeline event"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    event = TimelineEvent(
        case_id=case_id,
        user_id=user.user_id,
        **event_data.model_dump()
    )
    
    event_dict = event.model_dump()
    event_dict["event_date"] = event_dict["event_date"].isoformat()
    event_dict["created_at"] = event_dict["created_at"].isoformat()
    
    await db.timeline_events.insert_one(event_dict)
    
    # Update case
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Return the event without MongoDB's _id field
    created_event = await db.timeline_events.find_one({"event_id": event.event_id}, {"_id": 0})
    return created_event

@api_router.put("/cases/{case_id}/timeline/{event_id}", response_model=dict)
async def update_timeline_event(case_id: str, event_id: str, event_data: TimelineEventCreate, request: Request):
    """Update a timeline event"""
    user = await get_current_user(request)
    
    update_data = event_data.model_dump()
    update_data["event_date"] = update_data["event_date"].isoformat()
    
    result = await db.timeline_events.update_one(
        {"event_id": event_id, "case_id": case_id, "user_id": user.user_id},
        {"$set": update_data}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Event not found")
    
    return await db.timeline_events.find_one({"event_id": event_id}, {"_id": 0})

@api_router.delete("/cases/{case_id}/timeline/{event_id}")
async def delete_timeline_event(case_id: str, event_id: str, request: Request):
    """Delete a timeline event"""
    user = await get_current_user(request)
    
    result = await db.timeline_events.delete_one({
        "event_id": event_id,
        "case_id": case_id,
        "user_id": user.user_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Event not found")
    
    return {"message": "Event deleted"}

@api_router.post("/cases/{case_id}/timeline/auto-generate", response_model=dict)
async def auto_generate_timeline(case_id: str, request: Request):
    """AI-powered timeline generation from uploaded documents"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get documents with content
    documents = await db.documents.find(
        {"case_id": case_id, "content_text": {"$exists": True, "$ne": ""}},
        {"_id": 0, "file_data": 0}
    ).to_list(100)
    
    if not documents:
        raise HTTPException(status_code=400, detail="No documents with extracted text found. Please upload documents and extract text first.")
    
    # Build context from documents
    doc_context = f"CASE: {case.get('title', 'Unknown')}\nDEFENDANT: {case.get('defendant_name', 'Unknown')}\n\n"
    doc_context += "=== DOCUMENTS ===\n\n"
    
    for doc in documents:
        content = doc.get('content_text', '')[:3000]  # Limit per doc
        doc_context += f"DOCUMENT: {doc.get('filename')}\n"
        doc_context += f"CONTENT:\n{content}\n\n"
    
    system_prompt = """You are an expert legal analyst specializing in criminal cases. 
Your task is to extract a chronological timeline of events from case documents.

For each event, identify:
1. The DATE (in YYYY-MM-DD format if possible, or approximate like "2020-01" or "Early 2020")
2. A clear TITLE for the event (brief, descriptive)
3. A DESCRIPTION with relevant details
4. The EVENT TYPE: one of [incident, arrest, court_hearing, evidence, witness, legal_filing, verdict, appeal, other]

Return your response as a JSON array of events, ordered chronologically. Example:
[
  {"date": "2019-05-15", "title": "Initial Incident", "description": "The alleged incident occurred at...", "event_type": "incident"},
  {"date": "2019-05-16", "title": "Arrest of Defendant", "description": "Police arrested...", "event_type": "arrest"}
]

Only include events you can clearly identify from the documents. Be thorough - extract ALL dates and events mentioned."""

    user_prompt = f"""Analyze these documents and extract a complete chronological timeline of all events.

{doc_context}

Return ONLY a valid JSON array of timeline events. No other text."""

    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="AI service not configured")
    
    # Try up to 4 times with exponential backoff
    response = None
    last_error = None
    for attempt in range(4):
        try:
            chat = LlmChat(
                api_key=api_key,
                session_id=f"timeline_{case_id}_{uuid.uuid4().hex[:8]}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o")
            
            response = await chat.send_message(UserMessage(text=user_prompt))
            break
        except Exception as e:
            last_error = e
            logger.warning(f"Timeline generation attempt {attempt + 1} failed: {e}")
            if attempt < 3:
                import asyncio
                await asyncio.sleep(3 * (2 ** attempt))
    
    if response is None:
        logger.error(f"All timeline generation attempts failed: {last_error}")
        raise HTTPException(status_code=500, detail=f"AI timeline generation failed after retries: {str(last_error)}")
    
    # Parse the JSON response
    import json
    import re
    
    # Extract JSON array from response
    json_match = re.search(r'\[[\s\S]*\]', response)
    if not json_match:
        logger.warning(f"Could not parse timeline JSON from response: {response[:500]}")
        raise HTTPException(status_code=500, detail="Failed to parse timeline from AI response")
    
    try:
        events_data = json.loads(json_match.group())
    except json.JSONDecodeError as e:
        logger.warning(f"JSON decode error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse timeline JSON")
    
    # Create timeline events in database
    created_events = []
    for event in events_data:
        event_date = event.get('date', '')
        
        # Try to parse and normalize the date
        parsed_date = None
        if event_date:
            # Try various date formats
            for fmt in ['%Y-%m-%d', '%Y-%m', '%Y', '%d/%m/%Y', '%d-%m-%Y']:
                try:
                    parsed_date = datetime.strptime(event_date, fmt)
                    break
                except ValueError:
                    continue
            
            if not parsed_date:
                # Try to extract year at minimum
                year_match = re.search(r'(19|20)\d{2}', event_date)
                if year_match:
                    parsed_date = datetime.strptime(year_match.group(), '%Y')
        
        if not parsed_date:
            parsed_date = datetime.now(timezone.utc)
        
        # Validate event type
        valid_types = ['incident', 'arrest', 'court_hearing', 'evidence', 'witness', 'legal_filing', 'verdict', 'appeal', 'other']
        event_type = event.get('event_type', 'other')
        if event_type not in valid_types:
            event_type = 'other'
        
        timeline_event = {
            "event_id": f"evt_{uuid.uuid4().hex[:12]}",
            "case_id": case_id,
            "user_id": user.user_id,
            "title": event.get('title', 'Unknown Event')[:200],
            "description": event.get('description', '')[:2000],
            "event_date": parsed_date.isoformat(),
            "event_type": event_type,
            "auto_generated": True,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.timeline_events.insert_one(timeline_event)
        if '_id' in timeline_event:
            del timeline_event['_id']
        created_events.append(timeline_event)
    
    # Sort by date
    created_events.sort(key=lambda x: x.get('event_date', ''))
    
    return {
        "message": f"Successfully generated {len(created_events)} timeline events",
        "events_created": len(created_events),
        "events": created_events
    }

@api_router.post("/cases/{case_id}/timeline/analyze", response_model=dict)
async def analyze_timeline(case_id: str, request: Request):
    """AI-powered timeline analysis to find gaps, inconsistencies, and insights"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    events = await db.timeline_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("event_date", 1).to_list(500)
    
    if len(events) < 2:
        raise HTTPException(status_code=400, detail="Need at least 2 timeline events for analysis")
    
    # Get grounds for context
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    emergent_api_key = os.environ.get('EMERGENT_LLM_KEY')
    
    system_prompt = """You are an expert criminal appeals analyst specializing in NSW and Australian federal law.
Analyze this timeline of events and provide detailed insights for an appeal case.

Your analysis must include:
1. TIMELINE GAPS: Identify any suspicious gaps or missing periods that should have documentation
2. INCONSISTENCIES: Find contradictions between events or illogical sequences
3. PROSECUTION VS DEFENCE: Identify which events favor prosecution vs defence
4. CONTESTED FACTS: Flag events that appear disputed or have conflicting accounts
5. APPEAL RELEVANCE: Link events to potential grounds of appeal
6. KEY DATES: Highlight critical dates and any statute of limitation concerns
7. WITNESS PATTERNS: Note patterns in witness involvement across events

Return a JSON object with this structure:
{
    "gaps": [{"start_date": "...", "end_date": "...", "description": "...", "significance": "high/medium/low"}],
    "inconsistencies": [{"event_ids": [...], "description": "...", "impact": "..."}],
    "prosecution_events": [{"event_id": "...", "reason": "..."}],
    "defence_events": [{"event_id": "...", "reason": "..."}],
    "contested_facts": [{"event_id": "...", "issue": "...", "recommendation": "..."}],
    "ground_connections": [{"event_id": "...", "ground_type": "...", "relevance": "..."}],
    "key_observations": ["..."],
    "recommended_actions": ["..."]
}"""

    events_text = json.dumps(events, indent=2, default=str)
    grounds_text = json.dumps(grounds, indent=2, default=str) if grounds else "No grounds identified yet"
    
    user_prompt = f"""Analyze this criminal case timeline for an appeal:

TIMELINE EVENTS:
{events_text}

IDENTIFIED GROUNDS OF APPEAL:
{grounds_text}

Provide a comprehensive analysis identifying gaps, inconsistencies, and appeal-relevant insights."""

    last_error = None
    for attempt in range(3):
        try:
            chat = LlmChat(
                api_key=emergent_api_key,
                session_id=f"timeline_analysis_{case_id}_{uuid.uuid4().hex[:8]}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o")
            
            response = await chat.send_message(UserMessage(text=user_prompt))
            break
        except Exception as e:
            last_error = e
            if attempt < 2:
                await asyncio.sleep(2 ** attempt)
    else:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(last_error)}")
    
    # Parse JSON response
    try:
        json_match = response
        if "```json" in response:
            json_match = response.split("```json")[1].split("```")[0]
        elif "```" in response:
            json_match = response.split("```")[1].split("```")[0]
        
        analysis = json.loads(json_match.strip())
    except json.JSONDecodeError:
        # Return raw insights if JSON parsing fails
        analysis = {
            "gaps": [],
            "inconsistencies": [],
            "prosecution_events": [],
            "defence_events": [],
            "contested_facts": [],
            "ground_connections": [],
            "key_observations": [response[:2000]],
            "recommended_actions": []
        }
    
    return {
        "analysis": analysis,
        "event_count": len(events),
        "analyzed_at": datetime.now(timezone.utc).isoformat()
    }

@api_router.get("/cases/{case_id}/timeline/export-pdf")
async def export_timeline_pdf(case_id: str, request: Request):
    """Export timeline as a formatted PDF"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    events = await db.timeline_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("event_date", 1).to_list(500)
    
    # Get linked documents for reference
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0, "document_id": 1, "filename": 1}
    ).to_list(500)
    doc_map = {d["document_id"]: d["filename"] for d in documents}
    
    # Get grounds for reference
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0, "ground_id": 1, "title": 1}
    ).to_list(100)
    ground_map = {g["ground_id"]: g["title"] for g in grounds}
    
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from io import BytesIO
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm, leftMargin=15*mm, rightMargin=15*mm)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=12, textColor=colors.HexColor('#0f172a'))
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#64748b'), spaceAfter=20)
    event_title_style = ParagraphStyle('EventTitle', parent=styles['Heading3'], fontSize=12, textColor=colors.HexColor('#1e293b'), spaceBefore=8, spaceAfter=4)
    event_body_style = ParagraphStyle('EventBody', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#475569'), spaceAfter=4)
    meta_style = ParagraphStyle('Meta', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#94a3b8'))
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#334155'), spaceBefore=16, spaceAfter=8)
    
    story = []
    
    # Header
    story.append(Paragraph(f"Case Timeline: {case.get('title', 'Untitled Case')}", title_style))
    story.append(Paragraph(f"Generated for {user.name} on {datetime.now().strftime('%d %B %Y')}", subtitle_style))
    story.append(Paragraph("Criminal Appeal AI — Deb King, Glenmore Park 2745", meta_style))
    story.append(Spacer(1, 10*mm))
    
    # Summary stats
    critical_count = len([e for e in events if e.get('significance') == 'critical'])
    contested_count = len([e for e in events if e.get('is_contested')])
    story.append(Paragraph("Timeline Summary", section_style))
    summary_data = [
        ["Total Events", str(len(events))],
        ["Critical Events", str(critical_count)],
        ["Contested Facts", str(contested_count)],
        ["Date Range", f"{events[0].get('event_date', 'N/A')[:10] if events else 'N/A'} to {events[-1].get('event_date', 'N/A')[:10] if events else 'N/A'}"]
    ]
    summary_table = Table(summary_data, colWidths=[80*mm, 80*mm])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f1f5f9')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#334155')),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('PADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 8*mm))
    
    # Events
    story.append(Paragraph("Chronological Events", section_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
    
    significance_colors = {
        'critical': '#dc2626',
        'important': '#ea580c', 
        'normal': '#3b82f6',
        'minor': '#9ca3af'
    }
    
    for i, event in enumerate(events):
        date_str = event.get('event_date', '')[:10] if event.get('event_date') else 'Unknown'
        sig = event.get('significance', 'normal')
        sig_color = significance_colors.get(sig, '#3b82f6')
        
        contested_marker = " [CONTESTED]" if event.get('is_contested') else ""
        perspective = f" ({event.get('perspective', 'neutral').upper()})" if event.get('perspective') != 'neutral' else ""
        
        story.append(Paragraph(f"<font color='{sig_color}'>●</font> {date_str} — {event.get('title', 'Untitled')}{contested_marker}{perspective}", event_title_style))
        story.append(Paragraph(event.get('description', 'No description'), event_body_style))
        
        # Metadata
        meta_parts = []
        if event.get('event_type'):
            meta_parts.append(f"Type: {event['event_type'].replace('_', ' ').title()}")
        if event.get('source_citation'):
            meta_parts.append(f"Source: {event['source_citation'][:50]}")
        if event.get('linked_documents'):
            linked_names = [doc_map.get(d, d) for d in event['linked_documents'][:3]]
            meta_parts.append(f"Docs: {', '.join(linked_names)}")
        if event.get('related_grounds'):
            linked_grounds = [ground_map.get(g, g) for g in event['related_grounds'][:2]]
            meta_parts.append(f"Grounds: {', '.join(linked_grounds)}")
        if event.get('participants'):
            participants = [f"{p.get('name', 'Unknown')} ({p.get('role', 'Unknown')})" for p in event['participants'][:3]]
            meta_parts.append(f"Participants: {', '.join(participants)}")
        
        if meta_parts:
            story.append(Paragraph(" | ".join(meta_parts), meta_style))
        
        if event.get('contested_details'):
            story.append(Paragraph(f"<i>Contested: {event['contested_details']}</i>", meta_style))
        
        story.append(Spacer(1, 4*mm))
    
    # Footer
    story.append(Spacer(1, 10*mm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
    story.append(Paragraph("One woman's fight for justice — seeking truth for Joshua Homann, failed by the system", meta_style))
    
    doc.build(story)
    buffer.seek(0)
    
    filename = f"timeline_{case.get('title', 'case')[:30].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ============ DEADLINE ENDPOINTS ============

@api_router.get("/cases/{case_id}/deadlines", response_model=List[dict])
async def get_deadlines(case_id: str, request: Request):
    """Get all deadlines for a case"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    deadlines = await db.deadlines.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("due_date", 1).to_list(100)
    
    return deadlines

@api_router.post("/cases/{case_id}/deadlines", response_model=dict)
async def create_deadline(case_id: str, deadline_data: DeadlineCreate, request: Request):
    """Create a new deadline"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    deadline = Deadline(
        case_id=case_id,
        user_id=user.user_id,
        **deadline_data.model_dump()
    )
    
    deadline_dict = deadline.model_dump()
    deadline_dict["due_date"] = deadline_dict["due_date"].isoformat()
    deadline_dict["created_at"] = deadline_dict["created_at"].isoformat()
    
    await db.deadlines.insert_one(deadline_dict)
    
    return await db.deadlines.find_one({"deadline_id": deadline.deadline_id}, {"_id": 0})

@api_router.patch("/cases/{case_id}/deadlines/{deadline_id}", response_model=dict)
async def update_deadline(case_id: str, deadline_id: str, request: Request):
    """Update a deadline (mark complete, etc.)"""
    await get_current_user(request)  # Verify authentication
    body = await request.json()
    
    deadline = await db.deadlines.find_one({
        "deadline_id": deadline_id,
        "case_id": case_id
    })
    if not deadline:
        raise HTTPException(status_code=404, detail="Deadline not found")
    
    update_data = {}
    if "is_completed" in body:
        update_data["is_completed"] = body["is_completed"]
        if body["is_completed"]:
            update_data["completed_at"] = datetime.now(timezone.utc).isoformat()
    if "title" in body:
        update_data["title"] = body["title"]
    if "due_date" in body:
        update_data["due_date"] = body["due_date"]
    if "priority" in body:
        update_data["priority"] = body["priority"]
    
    if update_data:
        await db.deadlines.update_one(
            {"deadline_id": deadline_id},
            {"$set": update_data}
        )
    
    return await db.deadlines.find_one({"deadline_id": deadline_id}, {"_id": 0})

@api_router.delete("/cases/{case_id}/deadlines/{deadline_id}")
async def delete_deadline(case_id: str, deadline_id: str, request: Request):
    """Delete a deadline"""
    await get_current_user(request)  # Verify authentication
    
    result = await db.deadlines.delete_one({
        "deadline_id": deadline_id,
        "case_id": case_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Deadline not found")
    
    return {"message": "Deadline deleted"}

# ============ CHECKLIST ENDPOINTS ============

@api_router.get("/cases/{case_id}/checklist", response_model=List[dict])
async def get_checklist(case_id: str, request: Request):
    """Get checklist for a case, creating default items if none exist"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    items = await db.checklist_items.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("order", 1).to_list(100)
    
    if not items:
        for item_data in DEFAULT_CHECKLIST:
            item = ChecklistItem(
                case_id=case_id,
                user_id=user.user_id,
                **item_data
            )
            item_dict = item.model_dump()
            item_dict["created_at"] = item_dict["created_at"].isoformat()
            await db.checklist_items.insert_one(item_dict)
        
        items = await db.checklist_items.find(
            {"case_id": case_id},
            {"_id": 0}
        ).sort("order", 1).to_list(100)
    
    return items

@api_router.patch("/cases/{case_id}/checklist/{item_id}", response_model=dict)
async def update_checklist_item(case_id: str, item_id: str, request: Request):
    """Update a checklist item"""
    await get_current_user(request)  # Verify authentication
    body = await request.json()
    
    item = await db.checklist_items.find_one({
        "item_id": item_id,
        "case_id": case_id
    })
    if not item:
        raise HTTPException(status_code=404, detail="Checklist item not found")
    
    update_data = {}
    if "is_completed" in body:
        update_data["is_completed"] = body["is_completed"]
        if body["is_completed"]:
            update_data["completed_at"] = datetime.now(timezone.utc).isoformat()
        else:
            update_data["completed_at"] = None
    
    if update_data:
        await db.checklist_items.update_one(
            {"item_id": item_id},
            {"$set": update_data}
        )
    
    return await db.checklist_items.find_one({"item_id": item_id}, {"_id": 0})

# ============ CASE STRENGTH ENDPOINT ============

@api_router.get("/cases/{case_id}/strength", response_model=dict)
async def get_case_strength(case_id: str, request: Request):
    """Calculate overall case strength"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    grounds = await db.grounds_of_merit.find({"case_id": case_id}, {"_id": 0}).to_list(100)
    documents = await db.documents.find({"case_id": case_id}, {"_id": 0}).to_list(500)
    timeline = await db.timeline_events.find({"case_id": case_id}, {"_id": 0}).to_list(500)
    checklist = await db.checklist_items.find({"case_id": case_id}, {"_id": 0}).to_list(100)
    
    strength_scores = {"strong": 3, "moderate": 2, "weak": 1}
    
    grounds_score = 0
    ground_breakdown = {"strong": 0, "moderate": 0, "weak": 0}
    if grounds:
        for g in grounds:
            strength = g.get("strength", "moderate")
            ground_breakdown[strength] = ground_breakdown.get(strength, 0) + 1
            grounds_score += strength_scores.get(strength, 1)
        max_possible = len(grounds) * 3
        grounds_score = min(100, int((grounds_score / max_possible) * 100) + (len(grounds) * 5))
    
    doc_score = 0
    docs_with_text = len([d for d in documents if d.get("content_text")])
    if documents:
        doc_score = min(100, int((docs_with_text / len(documents)) * 50) + (len(documents) * 5))
    
    timeline_score = 0
    if timeline:
        timeline_score = min(100, len(timeline) * 5)
        critical_events = len([t for t in timeline if t.get("significance") == "critical"])
        timeline_score = min(100, timeline_score + (critical_events * 10))
    
    prep_score = 0
    if checklist:
        completed = len([c for c in checklist if c.get("is_completed")])
        prep_score = int((completed / len(checklist)) * 100)
    
    overall_score = int(
        (grounds_score * 0.40) +
        (doc_score * 0.25) +
        (timeline_score * 0.15) +
        (prep_score * 0.20)
    )
    
    if overall_score >= 75:
        rating, rating_color = "Strong", "green"
    elif overall_score >= 50:
        rating, rating_color = "Moderate", "amber"
    elif overall_score >= 25:
        rating, rating_color = "Developing", "orange"
    else:
        rating, rating_color = "Early Stage", "red"
    
    recommendations = []
    if not grounds:
        recommendations.append("Run AI Grounds Identification to find potential appeal grounds")
    if len(documents) < 3:
        recommendations.append("Upload more case documents for better analysis")
    if len(timeline) < 5:
        recommendations.append("Build out your timeline with key case events")
    if prep_score < 50:
        recommendations.append("Work through the appeal checklist")
    
    return {
        "overall_score": overall_score,
        "rating": rating,
        "rating_color": rating_color,
        "breakdown": {
            "grounds": {"score": grounds_score, "count": len(grounds), **ground_breakdown},
            "documentation": {"score": doc_score, "total_docs": len(documents), "with_text": docs_with_text},
            "timeline": {"score": timeline_score, "event_count": len(timeline)},
            "preparation": {"score": prep_score, "completed": len([c for c in checklist if c.get("is_completed")]), "total": len(checklist)}
        },
        "recommendations": recommendations
    }

# ============ CONTRADICTION FINDER ============

@api_router.post("/cases/{case_id}/analyze-contradictions", response_model=dict)
async def analyze_witness_contradictions(case_id: str, request: Request):
    """AI analysis to find contradictions in documents"""
    user = await get_current_user(request)
    
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    documents = await db.documents.find(
        {"case_id": case_id, "content_text": {"$exists": True, "$ne": ""}},
        {"_id": 0, "file_data": 0}
    ).to_list(100)
    
    if len(documents) < 2:
        raise HTTPException(status_code=400, detail="Need at least 2 documents with text to compare")
    
    doc_context = ""
    for i, doc in enumerate(documents):
        doc_context += f"\n=== DOCUMENT {i+1}: {doc.get('filename', 'Unknown')} ===\n"
        content = doc.get('content_text', '')[:4000]
        doc_context += f"Content:\n{content}\n"
    
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    emergent_api_key = os.environ.get('EMERGENT_LLM_KEY')
    
    system_prompt = """You are a legal analyst finding contradictions in witness statements. Find:
1. DIRECT CONTRADICTIONS - witnesses contradict each other
2. INTERNAL INCONSISTENCIES - witness contradicts themselves
3. TIMELINE CONFLICTS - times/dates don't align
4. FACTUAL DISCREPANCIES - differences in descriptions
5. OMISSIONS - facts missing from one account

Return JSON:
{
    "contradictions": [{"type": "...", "severity": "critical|significant|minor", "documents_involved": [], "description": "...", "appeal_relevance": "..."}],
    "summary": "...",
    "total_critical": 0, "total_significant": 0, "total_minor": 0
}"""

    last_error = None
    for attempt in range(3):
        try:
            chat = LlmChat(
                api_key=emergent_api_key,
                session_id=f"contradiction_{case_id}_{uuid.uuid4().hex[:8]}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o")
            
            response = await chat.send_message(UserMessage(text=f"Find contradictions in these documents:\n{doc_context}"))
            break
        except Exception as e:
            last_error = e
            if attempt < 2:
                await asyncio.sleep(2 ** attempt)
    else:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(last_error)}")
    
    try:
        json_match = response
        if "```json" in response:
            json_match = response.split("```json")[1].split("```")[0]
        analysis = json.loads(json_match.strip())
    except (json.JSONDecodeError, IndexError, ValueError):
        analysis = {"contradictions": [], "summary": response[:2000], "total_critical": 0, "total_significant": 0, "total_minor": 0}
    
    return {"analysis": analysis, "documents_analyzed": len(documents), "analyzed_at": datetime.now(timezone.utc).isoformat()}

# ============ PAYMENT ENDPOINTS ============

@api_router.get("/payments/prices")
async def get_feature_prices():
    """Get pricing for premium features"""
    return {
        "prices": FEATURE_PRICES,
        "currency": "AUD",
        "paypal_client_id": PAYPAL_CLIENT_ID if PAYPAL_CONFIGURED else "",
        "paypal_configured": PAYPAL_CONFIGURED
    }

@api_router.get("/cases/{case_id}/payments")
async def get_case_payments(case_id: str, request: Request):
    """Get all payments for a case"""
    user = await get_current_user(request)
    
    payments = await db.payments.find(
        {"case_id": case_id, "user_id": user.user_id, "status": "completed"},
        {"_id": 0}
    ).to_list(100)
    
    # Return which features are unlocked
    unlocked = {
        "grounds_of_merit": False,
        "full_report": False,
        "extensive_report": False
    }
    
    for payment in payments:
        if payment.get("feature_type") in unlocked:
            unlocked[payment["feature_type"]] = True
    
    return {
        "payments": payments,
        "unlocked_features": unlocked
    }

@api_router.post("/cases/{case_id}/payments/create-order")
async def create_payment_order(case_id: str, request: Request):
    """Create a PayPal order for a feature purchase"""
    if not PAYPAL_CONFIGURED:
        raise HTTPException(status_code=503, detail="Payment service not configured")
    
    user = await get_current_user(request)
    body = await request.json()
    feature_type = body.get("feature_type")
    
    if feature_type not in FEATURE_PRICES:
        raise HTTPException(status_code=400, detail="Invalid feature type")
    
    # Check if already purchased
    existing = await db.payments.find_one({
        "case_id": case_id,
        "user_id": user.user_id,
        "feature_type": feature_type,
        "status": "completed"
    })
    if existing:
        raise HTTPException(status_code=400, detail="Feature already purchased for this case")
    
    price_info = FEATURE_PRICES[feature_type]
    
    # Create PayPal order
    payment = paypalrestsdk.Payment({
        "intent": "sale",
        "payer": {"payment_method": "paypal"},
        "redirect_urls": {
            "return_url": f"{os.environ.get('FRONTEND_URL', 'http://localhost:3000')}/payment/success",
            "cancel_url": f"{os.environ.get('FRONTEND_URL', 'http://localhost:3000')}/payment/cancel"
        },
        "transactions": [{
            "item_list": {
                "items": [{
                    "name": price_info["name"],
                    "sku": feature_type,
                    "price": str(price_info["price"]),
                    "currency": "AUD",
                    "quantity": 1
                }]
            },
            "amount": {
                "total": str(price_info["price"]),
                "currency": "AUD"
            },
            "description": f"Criminal Appeal AI - {price_info['name']} for case {case_id}"
        }]
    })
    
    if payment.create():
        # Store pending payment
        payment_record = Payment(
            user_id=user.user_id,
            case_id=case_id,
            feature_type=feature_type,
            amount=price_info["price"],
            paypal_order_id=payment.id,
            status="pending"
        )
        payment_dict = payment_record.model_dump()
        payment_dict["created_at"] = payment_dict["created_at"].isoformat()
        await db.payments.insert_one(payment_dict)
        
        # Get approval URL
        approval_url = None
        for link in payment.links:
            if link.rel == "approval_url":
                approval_url = link.href
                break
        
        return {
            "payment_id": payment_record.payment_id,
            "paypal_order_id": payment.id,
            "approval_url": approval_url
        }
    else:
        logger.error(f"PayPal payment creation failed: {payment.error}")
        raise HTTPException(status_code=500, detail="Failed to create payment")

@api_router.post("/cases/{case_id}/payments/capture")
async def capture_payment(case_id: str, request: Request):
    """Capture/execute a PayPal payment after user approval"""
    user = await get_current_user(request)
    body = await request.json()
    paypal_payment_id = body.get("paypal_payment_id")
    payer_id = body.get("payer_id")
    
    if not paypal_payment_id or not payer_id:
        raise HTTPException(status_code=400, detail="paypal_payment_id and payer_id required")
    
    # Find the pending payment
    payment_record = await db.payments.find_one({
        "paypal_order_id": paypal_payment_id,
        "case_id": case_id,
        "user_id": user.user_id,
        "status": "pending"
    })
    
    if not payment_record:
        raise HTTPException(status_code=404, detail="Payment not found")
    
    # Execute the payment
    payment = paypalrestsdk.Payment.find(paypal_payment_id)
    
    if payment.execute({"payer_id": payer_id}):
        # Update payment record
        await db.payments.update_one(
            {"paypal_order_id": paypal_payment_id},
            {"$set": {
                "status": "completed",
                "completed_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        return {
            "success": True,
            "message": "Payment completed successfully",
            "feature_type": payment_record["feature_type"]
        }
    else:
        await db.payments.update_one(
            {"paypal_order_id": paypal_payment_id},
            {"$set": {"status": "failed"}}
        )
        raise HTTPException(status_code=400, detail="Payment execution failed")

@api_router.post("/payments/webhook")
async def paypal_webhook(request: Request):
    """Handle PayPal IPN/webhook notifications"""
    body = await request.json()
    event_type = body.get("event_type", "")
    
    if event_type == "PAYMENT.SALE.COMPLETED":
        resource = body.get("resource", {})
        parent_payment = resource.get("parent_payment")
        
        if parent_payment:
            await db.payments.update_one(
                {"paypal_order_id": parent_payment},
                {"$set": {
                    "status": "completed",
                    "completed_at": datetime.now(timezone.utc).isoformat()
                }}
            )
    
    return {"status": "received"}

# ============ RESOURCE DIRECTORY ============

@api_router.get("/resources/directory", response_model=dict)
async def get_resource_directory():
    """Get directory of support resources"""
    return {
        "support_services": [
            {"name": "Community Legal Centres NSW", "website": "https://www.clcnsw.org.au", "services": ["Legal advice", "Referrals"], "region": "NSW"},
            {"name": "Aboriginal Legal Service (NSW/ACT)", "phone": "1800 765 767", "website": "https://www.alsnswact.org.au", "services": ["Criminal law", "Family law"], "region": "NSW/ACT"},
            {"name": "LawAccess NSW", "phone": "1300 888 529", "website": "https://www.lawaccess.nsw.gov.au", "services": ["Legal information", "Referrals"], "region": "NSW"}
        ],
        "advocacy_groups": [
            {"name": "Innocence Project (Australia)", "website": "https://www.innocenceproject.org.au", "focus": "Wrongful convictions"},
            {"name": "Justice Action", "phone": "(02) 9283 0123", "website": "https://www.justiceaction.org.au", "focus": "Prisoner rights"},
            {"name": "Prisoners Aid Association NSW", "phone": "(02) 9288 8700", "website": "https://www.prisonersaid.org.au", "focus": "Family support"}
        ],
        "courts": [
            {"name": "NSW Court of Criminal Appeal", "website": "https://www.supremecourt.justice.nsw.gov.au"},
            {"name": "High Court of Australia", "website": "https://www.hcourt.gov.au", "note": "Special leave required"}
        ],
        "appeal_deadlines": {"notice_of_appeal": "28 days from conviction/sentence", "leave_to_appeal": "28 days", "extension": "Can apply if missed - must show good reason"}
    }

# ============ DOCUMENT TEMPLATES ============

@api_router.get("/templates", response_model=List[dict])
async def get_document_templates():
    """Get available document templates"""
    return [
        {"template_id": "notice_of_appeal", "title": "Notice of Appeal", "description": "Form to lodge an appeal", "category": "lodgement"},
        {"template_id": "leave_to_appeal", "title": "Application for Leave to Appeal", "description": "Application for permission to appeal sentence", "category": "lodgement"},
        {"template_id": "affidavit_fresh_evidence", "title": "Affidavit - Fresh Evidence", "description": "Sworn statement for new evidence", "category": "evidence"},
        {"template_id": "extension_of_time", "title": "Extension of Time Application", "description": "Apply to file after deadline", "category": "lodgement"},
        {"template_id": "outline_of_submissions", "title": "Written Submissions", "description": "Legal arguments for hearing", "category": "hearing"}
    ]

@api_router.post("/templates/{template_id}/generate", response_model=dict)
async def generate_document_from_template(template_id: str, request: Request):
    """Generate a document from template"""
    await get_current_user(request)  # Verify authentication
    body = await request.json()
    
    if template_id == "notice_of_appeal":
        content = f"""COURT OF CRIMINAL APPEAL - SUPREME COURT OF NSW
NOTICE OF APPEAL

Appellant: {body.get('appellant_name', '[NAME]')}
Case Number: {body.get('case_number', '[NUMBER]')}

TAKE NOTICE that the above-named appeals against: {body.get('appeal_type', '[CONVICTION/SENTENCE]')}

Court of Trial: {body.get('court_of_trial', '[COURT]')}
Date of Conviction: {body.get('date_of_conviction', '[DATE]')}
Date of Sentence: {body.get('date_of_sentence', '[DATE]')}
Offence: {body.get('offence', '[OFFENCE]')}
Sentence: {body.get('sentence', '[SENTENCE]')}

GROUNDS OF APPEAL:
{body.get('grounds', '[GROUNDS]')}

DATED: {datetime.now().strftime('%d %B %Y')}

---
Generated by Criminal Appeal AI - Deb King, Glenmore Park 2745"""
    elif template_id == "leave_to_appeal":
        content = f"""APPLICATION FOR LEAVE TO APPEAL AGAINST SENTENCE

Appellant: {body.get('appellant_name', '[NAME]')}
Case: {body.get('case_number', '[NUMBER]')}
Sentence Date: {body.get('date_of_sentence', '[DATE]')}
Sentence: {body.get('sentence', '[SENTENCE]')}

GROUNDS: {body.get('grounds', '[GROUNDS]')}
WHY LEAVE SHOULD BE GRANTED: {body.get('why_leave_granted', '[REASONS]')}

DATED: {datetime.now().strftime('%d %B %Y')}
---
Generated by Criminal Appeal AI"""
    else:
        content = f"Template {template_id} - Please contact support for this template."
    
    return {"template_id": template_id, "content": content, "generated_at": datetime.now(timezone.utc).isoformat()}

# ============ NOTES & COMMENTS ENDPOINTS ============

@api_router.get("/cases/{case_id}/notes", response_model=List[dict])
async def get_notes(case_id: str, request: Request):
    """Get all notes for a case"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    notes = await db.notes.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort([("is_pinned", -1), ("created_at", -1)]).to_list(500)
    
    return notes

@api_router.post("/cases/{case_id}/notes", response_model=dict)
async def create_note(case_id: str, note_data: NoteCreate, request: Request):
    """Create a new note"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    note = Note(
        case_id=case_id,
        user_id=user.user_id,
        author_name=user.name,
        author_email=user.email,
        **note_data.model_dump()
    )
    
    note_dict = note.model_dump()
    note_dict["created_at"] = note_dict["created_at"].isoformat()
    note_dict["updated_at"] = note_dict["updated_at"].isoformat()
    
    await db.notes.insert_one(note_dict)
    
    # Update case
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    created_note = await db.notes.find_one({"note_id": note.note_id}, {"_id": 0})
    return created_note

@api_router.get("/cases/{case_id}/notes/{note_id}", response_model=dict)
async def get_note(case_id: str, note_id: str, request: Request):
    """Get a specific note"""
    user = await get_current_user(request)
    
    note = await db.notes.find_one(
        {"note_id": note_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    return note

@api_router.put("/cases/{case_id}/notes/{note_id}", response_model=dict)
async def update_note(case_id: str, note_id: str, note_data: NoteUpdate, request: Request):
    """Update a note"""
    user = await get_current_user(request)
    
    update_fields = {k: v for k, v in note_data.model_dump().items() if v is not None}
    update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.notes.update_one(
        {"note_id": note_id, "case_id": case_id, "user_id": user.user_id},
        {"$set": update_fields}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")
    
    return await db.notes.find_one({"note_id": note_id}, {"_id": 0})

@api_router.delete("/cases/{case_id}/notes/{note_id}")
async def delete_note(case_id: str, note_id: str, request: Request):
    """Delete a note"""
    user = await get_current_user(request)
    
    result = await db.notes.delete_one({
        "note_id": note_id,
        "case_id": case_id,
        "user_id": user.user_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Note not found")
    
    return {"message": "Note deleted"}

@api_router.patch("/cases/{case_id}/notes/{note_id}/pin")
async def toggle_pin_note(case_id: str, note_id: str, request: Request):
    """Toggle pin status of a note"""
    user = await get_current_user(request)
    
    note = await db.notes.find_one(
        {"note_id": note_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    new_pin_status = not note.get("is_pinned", False)
    
    await db.notes.update_one(
        {"note_id": note_id},
        {"$set": {"is_pinned": new_pin_status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return await db.notes.find_one({"note_id": note_id}, {"_id": 0})

# ============ GROUNDS OF MERIT ENDPOINTS ============

GROUND_TYPES = [
    "procedural_error",
    "fresh_evidence", 
    "miscarriage_of_justice",
    "sentencing_error",
    "judicial_error",
    "ineffective_counsel",
    "prosecution_misconduct",
    "jury_irregularity",
    "constitutional_violation",
    "other"
]

@api_router.get("/cases/{case_id}/grounds", response_model=dict)
async def get_grounds_of_merit(case_id: str, request: Request):
    """Get all grounds of merit for a case - requires payment to see details"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort([("strength", 1), ("created_at", -1)]).to_list(100)
    
    # Check if user has paid for grounds access
    payment = await db.payments.find_one({
        "case_id": case_id,
        "user_id": user.user_id,
        "feature_type": "grounds_of_merit",
        "status": "completed"
    })
    
    is_unlocked = payment is not None
    
    if is_unlocked:
        # Return full grounds data
        return {
            "grounds": grounds,
            "count": len(grounds),
            "is_unlocked": True,
            "unlock_price": FEATURE_PRICES["grounds_of_merit"]["price"]
        }
    else:
        # Return only count and basic info (titles only, no descriptions)
        preview_grounds = []
        for g in grounds:
            preview_grounds.append({
                "ground_id": g.get("ground_id"),
                "title": g.get("title"),
                "ground_type": g.get("ground_type"),
                "strength": g.get("strength"),
                # Hide detailed content
                "description": "*** UNLOCK TO VIEW ***",
                "supporting_evidence": [],
                "analysis": None
            })
        return {
            "grounds": preview_grounds,
            "count": len(grounds),
            "is_unlocked": False,
            "unlock_price": FEATURE_PRICES["grounds_of_merit"]["price"],
            "message": f"Found {len(grounds)} potential grounds of merit. Pay ${FEATURE_PRICES['grounds_of_merit']['price']:.2f} to unlock full details."
        }

@api_router.post("/cases/{case_id}/grounds", response_model=dict)
async def create_ground_of_merit(case_id: str, ground_data: GroundOfMeritCreate, request: Request):
    """Create a new ground of merit"""
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    ground = GroundOfMerit(
        case_id=case_id,
        user_id=user.user_id,
        **ground_data.model_dump()
    )
    
    ground_dict = ground.model_dump()
    ground_dict["created_at"] = ground_dict["created_at"].isoformat()
    ground_dict["updated_at"] = ground_dict["updated_at"].isoformat()
    
    await db.grounds_of_merit.insert_one(ground_dict)
    
    # Update case
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    created_ground = await db.grounds_of_merit.find_one({"ground_id": ground.ground_id}, {"_id": 0})
    return created_ground

@api_router.get("/cases/{case_id}/grounds/{ground_id}", response_model=dict)
async def get_ground_of_merit(case_id: str, ground_id: str, request: Request):
    """Get a specific ground of merit"""
    user = await get_current_user(request)
    
    ground = await db.grounds_of_merit.find_one(
        {"ground_id": ground_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not ground:
        raise HTTPException(status_code=404, detail="Ground of merit not found")
    
    return ground

@api_router.put("/cases/{case_id}/grounds/{ground_id}", response_model=dict)
async def update_ground_of_merit(case_id: str, ground_id: str, ground_data: GroundOfMeritUpdate, request: Request):
    """Update a ground of merit"""
    user = await get_current_user(request)
    
    update_fields = {k: v for k, v in ground_data.model_dump().items() if v is not None}
    update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.grounds_of_merit.update_one(
        {"ground_id": ground_id, "case_id": case_id, "user_id": user.user_id},
        {"$set": update_fields}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ground of merit not found")
    
    return await db.grounds_of_merit.find_one({"ground_id": ground_id}, {"_id": 0})

@api_router.delete("/cases/{case_id}/grounds/{ground_id}")
async def delete_ground_of_merit(case_id: str, ground_id: str, request: Request):
    """Delete a ground of merit"""
    user = await get_current_user(request)
    
    result = await db.grounds_of_merit.delete_one({
        "ground_id": ground_id,
        "case_id": case_id,
        "user_id": user.user_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ground of merit not found")
    
    return {"message": "Ground of merit deleted"}

@api_router.post("/cases/{case_id}/grounds/{ground_id}/investigate", response_model=dict)
async def investigate_ground_of_merit(case_id: str, ground_id: str, request: Request):
    """Deep AI investigation of a specific ground of merit"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    user = await get_current_user(request)
    
    # Get the ground
    ground = await db.grounds_of_merit.find_one(
        {"ground_id": ground_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not ground:
        raise HTTPException(status_code=404, detail="Ground of merit not found")
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    
    # Get ALL documents with FULL content for cross-referencing
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0, "file_data": 0}
    ).to_list(500)
    
    # Get timeline
    timeline = await db.timeline_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("event_date", 1).to_list(500)
    
    # Build comprehensive context with FULL document content
    context = f"""
=== CASE INFORMATION ===
Title: {case.get('title', 'Unknown')}
Defendant: {case.get('defendant_name', 'Unknown')}
Case Number: {case.get('case_number', 'N/A')}
Court: {case.get('court', 'N/A')}

=== GROUND OF MERIT TO INVESTIGATE ===
Title: {ground.get('title')}
Type: {ground.get('ground_type')}
Description: {ground.get('description')}
Current Strength: {ground.get('strength')}
Supporting Evidence Listed: {', '.join(ground.get('supporting_evidence', []))}

"""
    
    # Include FULL document content
    if documents:
        context += f"=== CASE DOCUMENTS ({len(documents)} files) - SEARCH THESE FOR EVIDENCE ===\n"
        for doc in documents:
            context += f"\n--- DOCUMENT: {doc.get('filename')} [{doc.get('category')}] ---\n"
            content = doc.get('content_text', '')
            if content:
                context += f"FULL CONTENT:\n{content[:5000]}\n"
                if len(content) > 5000:
                    context += f"[... {len(content) - 5000} more characters ...]\n"
            else:
                context += "[No text content]\n"
    
    if timeline:
        context += f"\n=== TIMELINE ({len(timeline)} events) ===\n"
        for event in timeline:
            context += f"- {event.get('event_date')}: {event.get('title')} - {event.get('description', '')[:200]}\n"

    system_prompt = """You are an Australian criminal appeal barrister expert in NSW and Federal murder law.
IMPORTANT: You must search through ALL the document content provided and cite specific evidence.
Quote directly from documents to support your analysis."""

    user_prompt = f"""Conduct a THOROUGH investigation of this ground of merit.

{context}

REQUIRED ANALYSIS (search the documents above for evidence):

1. VIABILITY ASSESSMENT
   - Rate: Strong/Moderate/Weak with detailed justification
   - Quote specific evidence FROM THE DOCUMENTS above

2. DOCUMENT EVIDENCE
   - For EACH document, explain what it contains relevant to this ground
   - Quote specific passages that support or undermine this ground

3. RELEVANT LAW SECTIONS (be specific)
   - s.18 Crimes Act 1900 (NSW) - Murder
   - s.23 Crimes Act 1900 (NSW) - Provocation
   - s.6 Criminal Appeal Act 1912 (NSW) - Grounds for appeal
   - Other relevant sections with explanations

4. SIMILAR CASES
   - Name 2-3 Australian cases with citations
   - Explain relevance to this ground

5. EVIDENCE GAPS
   - What additional evidence would strengthen this ground?

6. STRATEGIC RECOMMENDATION
   - How to present this ground to the court"""

    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="AI service not configured")
    
    # Try up to 4 times with exponential backoff
    response = None
    last_error = None
    for attempt in range(4):
        try:
            chat = LlmChat(
                api_key=api_key,
                session_id=f"ground_{ground_id}_{uuid.uuid4().hex[:8]}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o")
            
            response = await chat.send_message(UserMessage(text=user_prompt))
            break
        except Exception as e:
            last_error = e
            logger.warning(f"Ground investigation attempt {attempt + 1} failed: {e}")
            if attempt < 3:
                import asyncio
                await asyncio.sleep(3 * (2 ** attempt))
    
    if response is None:
        logger.error(f"All ground investigation attempts failed: {last_error}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed after retries: {str(last_error)}")
    
    # Parse response to extract structured data
    law_sections = []
    similar_cases = []
    
    # Extract law sections
    import re
    section_patterns = re.findall(r'[sS]\.?\s*(\d+[A-Za-z]?)\s+([A-Za-z\s]+(?:Act|Code))\s*(?:\d{4})?', response)
    for section_num, act_name in section_patterns[:10]:
        law_sections.append({
            "section": section_num,
            "act": act_name.strip(),
            "jurisdiction": "NSW" if "NSW" in act_name or "1900" in response else "Federal"
        })
    
    # Extract case citations
    case_patterns = re.findall(r'([A-Z][a-z]+(?:\s+v\s+)[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', response)
    for case_name in list(set(case_patterns))[:5]:
        similar_cases.append({
            "case_name": case_name,
            "year": "",
            "citation": ""
        })
    
    # Update the ground with investigation results
    deep_analysis = {
        "full_analysis": response,
        "investigated_at": datetime.now(timezone.utc).isoformat(),
        "law_sections_identified": len(law_sections),
        "similar_cases_found": len(similar_cases),
        "documents_analyzed": len(documents)
    }
    
    await db.grounds_of_merit.update_one(
        {"ground_id": ground_id},
        {"$set": {
            "status": "investigated",
            "analysis": response[:2000] + "..." if len(response) > 2000 else response,
            "deep_analysis": deep_analysis,
            "law_sections": law_sections if law_sections else ground.get("law_sections", []),
            "similar_cases": similar_cases if similar_cases else ground.get("similar_cases", []),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return await db.grounds_of_merit.find_one({"ground_id": ground_id}, {"_id": 0})

@api_router.post("/cases/{case_id}/grounds/auto-identify", response_model=dict)
async def auto_identify_grounds(case_id: str, request: Request):
    """AI automatically identifies potential grounds of merit from case materials - prevents duplicates"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    user = await get_current_user(request)
    
    # Verify case ownership
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get existing grounds to avoid duplicates
    existing_grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Get all case materials - include content_text for analysis
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0, "file_data": 0}
    ).to_list(500)
    
    timeline = await db.timeline_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("event_date", 1).to_list(500)
    
    notes = await db.notes.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(500)
    
    # Build comprehensive context with full document content
    context = f"""
CRIMINAL APPEAL CASE ANALYSIS

CASE DETAILS:
- Title: {case.get('title', 'Unknown')}
- Defendant: {case.get('defendant_name', 'Unknown')}
- Case Number: {case.get('case_number', 'N/A')}
- Court: {case.get('court', 'N/A')}
- Judge: {case.get('judge', 'N/A')}
- Summary: {case.get('summary', 'No summary provided')}

"""
    
    # Include existing grounds so AI doesn't duplicate them
    if existing_grounds:
        context += f"=== ALREADY IDENTIFIED GROUNDS ({len(existing_grounds)}) ===\n"
        context += "DO NOT re-identify these grounds. Only identify NEW grounds not listed here:\n\n"
        for g in existing_grounds:
            context += f"- [{g.get('ground_type')}] {g.get('title')}\n"
            context += f"  Status: {g.get('status')}, Strength: {g.get('strength')}\n"
        context += "\n"
    
    # Include substantial document content for AI to analyze
    if documents:
        context += f"=== CASE DOCUMENTS ({len(documents)} files) ===\n\n"
        for doc in documents:
            context += f"--- DOCUMENT: {doc.get('filename')} ---\n"
            context += f"Category: {doc.get('category', 'other')}\n"
            if doc.get('description'):
                context += f"Description: {doc.get('description')}\n"
            
            # Include up to 5000 chars of content per document for thorough analysis
            content = doc.get('content_text', '')
            if content:
                context += f"CONTENT:\n{content[:5000]}\n"
                if len(content) > 5000:
                    context += f"[... {len(content) - 5000} more characters ...]\n"
            else:
                context += "CONTENT: [No text content extracted - may be image/scan]\n"
            context += "\n"
    else:
        context += "NO DOCUMENTS UPLOADED YET - Analysis based on case summary only.\n\n"

    if timeline:
        context += f"=== TIMELINE OF EVENTS ({len(timeline)} events) ===\n"
        for event in timeline:
            context += f"- {event.get('event_date', 'Unknown date')} [{event.get('event_type', 'event')}]: {event.get('title')}\n"
            if event.get('description'):
                context += f"  Details: {event.get('description')}\n"
        context += "\n"

    if notes:
        context += f"=== LEGAL NOTES & OBSERVATIONS ({len(notes)} notes) ===\n"
        for note in notes:
            context += f"- [{note.get('category', 'general')}] {note.get('title')}:\n"
            context += f"  {note.get('content', '')[:500]}\n"
        context += "\n"

    system_prompt = """You are a senior Australian criminal appeal barrister with 30+ years experience in murder and serious criminal appeals in NSW. You have successfully overturned dozens of wrongful convictions.

YOUR TASK: Conduct an EXHAUSTIVE, METICULOUS analysis of ALL provided case documents to identify EVERY possible ground of appeal. Leave no stone unturned.

ANALYSIS APPROACH - You MUST:
1. READ EVERY DOCUMENT THOROUGHLY - Extract all relevant facts, dates, witness statements, and evidence
2. CROSS-REFERENCE between documents - Look for contradictions, inconsistencies, and gaps
3. IDENTIFY PROCEDURAL ISSUES - Did police/prosecution follow correct procedures?
4. EXAMINE EVIDENCE HANDLING - Was evidence properly obtained, stored, disclosed?
5. REVIEW WITNESS TESTIMONY - Look for unreliable identification, coached witnesses, inconsistent statements
6. ASSESS LEGAL REPRESENTATION - Was the defence counsel competent? Did they call all witnesses? Cross-examine effectively?
7. CHECK JUDICIAL CONDUCT - Were jury directions proper? Was the judge impartial?
8. LOOK FOR FRESH EVIDENCE - Any new information that wasn't available at trial?
9. EXAMINE EXPERT EVIDENCE - Was forensic/medical evidence properly challenged?
10. CONSIDER DISCLOSURE FAILURES - Did prosecution disclose all relevant material?

GROUND TYPES TO IDENTIFY:
- PROCEDURAL ERROR: Police/court procedure violations, warrant issues, PACE breaches
- FRESH EVIDENCE: New witnesses, DNA evidence, alibis, recanted testimony
- MISCARRIAGE OF JUSTICE: Wrongful conviction indicators, unsafe verdict
- JUDICIAL ERROR: Wrong directions to jury, biased conduct, evidentiary rulings
- INEFFECTIVE COUNSEL: Failure to call witnesses, poor cross-examination, missed defences
- PROSECUTION MISCONDUCT: Non-disclosure, improper statements, witness coaching
- JURY IRREGULARITY: Misconduct, bias, improper deliberations
- SENTENCING ERROR: Manifestly excessive, wrong principles applied
- CONSTITUTIONAL/RIGHTS VIOLATIONS: Unfair trial, self-incrimination issues

BE THOROUGH - It is better to identify 10 potential grounds than miss 1 valid one. The appellant's liberty depends on this analysis.

IMPORTANT: If grounds are listed in "ALREADY IDENTIFIED GROUNDS", do NOT duplicate them. Only add NEW grounds."""

    user_prompt = f"""CONDUCT A COMPREHENSIVE LEGAL ANALYSIS OF THIS CRIMINAL APPEAL CASE:

{context}

INSTRUCTIONS:
1. Analyze EVERY document provided - read the full content carefully
2. Compare witness statements for inconsistencies
3. Identify ANY procedural irregularities
4. Look for evidence that may have been improperly handled or excluded
5. Consider whether the defence was adequate
6. Check for any fresh evidence possibilities
7. Assess whether the verdict was safe

For EACH ground you identify, provide DETAILED analysis including:
- Specific references to documents/evidence
- Why this constitutes a valid ground of appeal
- Relevant NSW and Federal law sections
- Assessment of strength (strong/moderate/weak)
- What evidence supports this ground

Return your analysis in this JSON format:
{{
  "grounds": [
    {{
      "title": "Specific, descriptive title",
      "ground_type": "procedural_error|fresh_evidence|miscarriage_of_justice|sentencing_error|judicial_error|ineffective_counsel|prosecution_misconduct|jury_irregularity|constitutional_violation|other",
      "description": "DETAILED description (at least 3-4 sentences) explaining the ground, citing specific evidence from the documents",
      "strength": "strong|moderate|weak",
      "key_evidence": ["Specific document references and quotes that support this ground"],
      "relevant_law": ["Specific law sections e.g., 's.18 Crimes Act 1900 (NSW)', 'Evidence Act 1995 s.137'"]
    }}
  ],
  "summary": "Overall assessment of appeal prospects and most promising grounds",
  "analysis_notes": "Any additional observations about the case"
}}

BE THOROUGH. Identify ALL potential grounds. The appellant's freedom may depend on this analysis."""

    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="AI service not configured")
    
    # Try up to 4 times with exponential backoff
    response = None
    last_error = None
    for attempt in range(4):
        try:
            chat = LlmChat(
                api_key=api_key,
                session_id=f"auto_identify_{case_id}_{uuid.uuid4().hex[:8]}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o")
            
            response = await chat.send_message(UserMessage(text=user_prompt))
            break
        except Exception as e:
            last_error = e
            logger.warning(f"Auto-identify attempt {attempt + 1} failed: {e}")
            if attempt < 3:
                import asyncio
                # Exponential backoff: 3, 6, 12 seconds
                await asyncio.sleep(3 * (2 ** attempt))
    
    if response is None:
        logger.error(f"All auto-identify attempts failed: {last_error}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed after retries: {str(last_error)}")
    
    # Helper function to check if ground is a duplicate
    def is_duplicate_ground(new_title, new_type, existing_grounds):
        new_title_lower = new_title.lower().strip()
        new_type_lower = new_type.lower().strip()
        
        for existing in existing_grounds:
            existing_title = existing.get('title', '').lower().strip()
            existing_type = existing.get('ground_type', '').lower().strip()
            
            # Same type and similar title = duplicate
            if existing_type == new_type_lower:
                # Check for significant title overlap
                new_words = set(new_title_lower.split())
                existing_words = set(existing_title.split())
                common_words = new_words & existing_words
                # Remove common legal words from comparison
                stop_words = {'the', 'a', 'an', 'of', 'in', 'to', 'for', 'and', 'or', 'at', 'by', 'on', 'with'}
                common_meaningful = common_words - stop_words
                new_meaningful = new_words - stop_words
                
                if len(new_meaningful) > 0:
                    overlap_ratio = len(common_meaningful) / len(new_meaningful)
                    if overlap_ratio > 0.5:  # More than 50% overlap
                        return True
            
            # Very similar titles regardless of type
            if new_title_lower == existing_title:
                return True
        
        return False
    
    # Try to parse JSON from response
    identified_grounds = []
    skipped_duplicates = 0
    try:
        import re
        json_match = re.search(r'\{[\s\S]*"grounds"[\s\S]*\}', response)
        if json_match:
            parsed = json.loads(json_match.group())
            grounds_data = parsed.get("grounds", [])
            
            for g in grounds_data:
                new_title = g.get("title", "Identified Ground")
                new_type = g.get("ground_type", "other")
                
                # Skip if duplicate
                if is_duplicate_ground(new_title, new_type, existing_grounds + identified_grounds):
                    skipped_duplicates += 1
                    logger.info(f"Skipping duplicate ground: {new_title}")
                    continue
                
                ground = GroundOfMerit(
                    case_id=case_id,
                    user_id=user.user_id,
                    title=new_title,
                    ground_type=new_type,
                    description=g.get("description", ""),
                    strength=g.get("strength", "moderate"),
                    supporting_evidence=g.get("key_evidence", []),
                    status="identified"
                )
                
                ground_dict = ground.model_dump()
                ground_dict["created_at"] = ground_dict["created_at"].isoformat()
                ground_dict["updated_at"] = ground_dict["updated_at"].isoformat()
                
                await db.grounds_of_merit.insert_one(ground_dict)
                # Fetch the inserted ground without _id
                created_ground = await db.grounds_of_merit.find_one({"ground_id": ground.ground_id}, {"_id": 0})
                identified_grounds.append(created_ground)
    except Exception as e:
        logger.error(f"Failed to parse AI response: {e}")
        # Only create fallback ground if no grounds were identified
        if len(identified_grounds) == 0:
            ground = GroundOfMerit(
                case_id=case_id,
                user_id=user.user_id,
                title="AI Analysis Results",
                ground_type="other",
                description="See analysis for identified grounds",
                strength="moderate",
                analysis=response,
                status="identified"
            )
            ground_dict = ground.model_dump()
            ground_dict["created_at"] = ground_dict["created_at"].isoformat()
            ground_dict["updated_at"] = ground_dict["updated_at"].isoformat()
            await db.grounds_of_merit.insert_one(ground_dict)
            # Fetch the inserted ground without _id
            created_ground = await db.grounds_of_merit.find_one({"ground_id": ground.ground_id}, {"_id": 0})
            identified_grounds.append(created_ground)
    
    # Update case
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return {
        "identified_count": len(identified_grounds),
        "skipped_duplicates": skipped_duplicates,
        "existing_grounds": len(existing_grounds),
        "grounds": identified_grounds,
        "message": f"Found {len(identified_grounds)} new grounds. Skipped {skipped_duplicates} duplicates." if skipped_duplicates > 0 else f"Found {len(identified_grounds)} new grounds."
    }

# ============ AI ANALYSIS & REPORTS ============

async def analyze_case_with_ai(case_id: str, user_id: str, report_type: str) -> dict:
    """Use OpenAI GPT-5.2 to analyze case and generate report"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get documents with full content
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0, "file_data": 0}
    ).to_list(500)
    
    # Get timeline
    timeline = await db.timeline_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("event_date", 1).to_list(500)
    
    # Get notes
    notes = await db.notes.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Get identified grounds
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Prepare comprehensive context for AI with FULL document content
    case_context = f"""
=== CASE INFORMATION ===
Title: {case.get('title', 'N/A')}
Defendant: {case.get('defendant_name', 'N/A')}
Case Number: {case.get('case_number', 'N/A')}
Court: {case.get('court', 'N/A')}
Judge: {case.get('judge', 'N/A')}
Summary: {case.get('summary', 'N/A')}

"""
    
    # Include FULL document content for cross-referencing
    if documents:
        case_context += f"=== CASE DOCUMENTS ({len(documents)} files) ===\n"
        for doc in documents:
            case_context += f"\n--- DOCUMENT: {doc.get('filename')} [{doc.get('category', 'other')}] ---\n"
            if doc.get('description'):
                case_context += f"Description: {doc.get('description')}\n"
            content = doc.get('content_text', '')
            if content:
                # Limit content per document based on report type to avoid timeouts
                max_chars = 2000 if report_type == "quick_summary" else 3000
                case_context += f"CONTENT:\n{content[:max_chars]}\n"
                if len(content) > max_chars:
                    case_context += f"[... {len(content) - max_chars} more characters ...]\n"
            else:
                case_context += "CONTENT: [No text extracted]\n"
    else:
        case_context += "=== NO DOCUMENTS UPLOADED ===\n"

    if timeline:
        case_context += f"\n=== TIMELINE OF EVENTS ({len(timeline)} events) ===\n"
        for event in timeline:
            case_context += f"- {event.get('event_date', 'Unknown')}: [{event.get('event_type')}] {event.get('title')}\n"
            if event.get('description'):
                case_context += f"  Details: {event.get('description')}\n"

    if notes:
        case_context += f"\n=== LEGAL NOTES ({len(notes)} notes) ===\n"
        for note in notes:
            case_context += f"- [{note.get('category')}] {note.get('title')}: {note.get('content', '')[:500]}\n"

    if grounds:
        case_context += f"\n=== IDENTIFIED GROUNDS OF MERIT ({len(grounds)} grounds) ===\n"
        for g in grounds:
            case_context += f"- [{g.get('ground_type')}] {g.get('title')} (Strength: {g.get('strength')})\n"
            case_context += f"  {g.get('description', '')[:300]}\n"

    # Define prompts based on report type
    if report_type == "quick_summary":
        system_prompt = """You are an expert criminal appeal legal analyst specializing in NSW State and Australian Federal murder law. 
Provide a concise summary of the case highlighting key points for an appeal. Reference specific documents and evidence."""
        user_prompt = f"""Analyze this criminal appeal case and provide a QUICK SUMMARY (2-3 paragraphs) including:
1. Brief case overview
2. Key evidence points (cite specific documents)
3. Primary grounds for appeal consideration

IMPORTANT: Cross-reference the actual document content provided below.

{case_context}"""

    elif report_type == "full_detailed":
        system_prompt = """You are an expert criminal appeal legal analyst for NSW and Australian Federal murder law. Cite specific documents and law sections."""
        user_prompt = f"""Provide a FULL DETAILED REPORT on this appeal case.

{case_context}

Include:
1. CASE OVERVIEW - cite specific documents
2. GROUNDS OF MERIT - with evidence quotes and law sections (NSW Crimes Act, Criminal Appeal Act, Criminal Code Act)
3. LEGAL FRAMEWORK - relevant law sections
4. RECOMMENDATIONS

Format as a legal brief for a barrister."""

    else:  # extensive_log
        system_prompt = """You are an expert criminal appeal legal analyst. Provide thorough documentation with document citations."""
        user_prompt = f"""Create an EXTENSIVE LOG REPORT for this case.

{case_context}

Include:
1. CASE CHRONOLOGY - with document citations
2. EVIDENCE ANALYSIS - quote key passages
3. ALL GROUNDS OF MERIT - with supporting evidence
4. LEGAL FRAMEWORK - with section numbers
5. DETAILED RECOMMENDATIONS"""

    # Call OpenAI via Emergent
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="AI service not configured")
    
    # Try up to 4 times with exponential backoff
    response = None
    last_error = None
    for attempt in range(4):
        try:
            chat = LlmChat(
                api_key=api_key,
                session_id=f"report_{case_id}_{uuid.uuid4().hex[:8]}",
                system_message=system_prompt
            ).with_model("openai", "gpt-4o")
            
            response = await chat.send_message(UserMessage(text=user_prompt))
            break
        except Exception as e:
            last_error = e
            logger.warning(f"Report generation attempt {attempt + 1} failed: {e}")
            if attempt < 3:
                import asyncio
                # Exponential backoff: 3, 6, 12 seconds
                await asyncio.sleep(3 * (2 ** attempt))
    
    if response is None:
        logger.error(f"All report generation attempts failed: {last_error}")
        raise HTTPException(status_code=500, detail=f"AI report generation failed after retries: {str(last_error)}")
    
    # Parse response to extract grounds of merit
    grounds_of_merit = []
    if "GROUNDS OF MERIT" in response or "Ground" in response:
        grounds_of_merit = [{
            "title": "AI-Identified Ground",
            "description": "See full report for details",
            "strength": "To be assessed by legal professional"
        }]
    
    return {
        "analysis": response,
        "grounds_of_merit": grounds_of_merit,
        "case_data": case,
        "document_count": len(documents),
        "event_count": len(timeline)
    }

@api_router.post("/cases/{case_id}/reports/generate", response_model=dict)
async def generate_report(case_id: str, report_request: ReportRequest, request: Request):
    """Generate an AI-powered report for a case"""
    user = await get_current_user(request)
    report_type = report_request.report_type
    
    if report_type not in ["quick_summary", "full_detailed", "extensive_log"]:
        raise HTTPException(status_code=400, detail="Invalid report type")
    
    # Check payment for premium reports
    if report_type == "full_detailed":
        payment = await db.payments.find_one({
            "case_id": case_id,
            "user_id": user.user_id,
            "feature_type": "full_report",
            "status": "completed"
        })
        if not payment:
            raise HTTPException(
                status_code=402, 
                detail={
                    "message": "Payment required for Full Detailed Report",
                    "feature_type": "full_report",
                    "price": FEATURE_PRICES["full_report"]["price"],
                    "currency": "AUD"
                }
            )
    
    if report_type == "extensive_log":
        payment = await db.payments.find_one({
            "case_id": case_id,
            "user_id": user.user_id,
            "feature_type": "extensive_report",
            "status": "completed"
        })
        if not payment:
            raise HTTPException(
                status_code=402, 
                detail={
                    "message": "Payment required for Extensive Log Report",
                    "feature_type": "extensive_report",
                    "price": FEATURE_PRICES["extensive_report"]["price"],
                    "currency": "AUD"
                }
            )
    
    # Generate AI analysis
    analysis_result = await analyze_case_with_ai(case_id, user.user_id, report_type)
    
    # Create report
    report_titles = {
        "quick_summary": "Quick Case Summary",
        "full_detailed": "Full Detailed Legal Analysis",
        "extensive_log": "Extensive Case Log & Analysis"
    }
    
    report = Report(
        case_id=case_id,
        user_id=user.user_id,
        report_type=report_type,
        title=report_titles[report_type],
        content={
            "analysis": analysis_result["analysis"],
            "case_title": analysis_result["case_data"].get("title", ""),
            "defendant": analysis_result["case_data"].get("defendant_name", ""),
            "document_count": analysis_result["document_count"],
            "event_count": analysis_result["event_count"]
        },
        grounds_of_merit=analysis_result["grounds_of_merit"]
    )
    
    report_dict = report.model_dump()
    report_dict["generated_at"] = report_dict["generated_at"].isoformat()
    
    await db.reports.insert_one(report_dict)
    
    # Return the report without MongoDB's _id field
    created_report = await db.reports.find_one({"report_id": report.report_id}, {"_id": 0})
    return created_report

@api_router.get("/cases/{case_id}/reports", response_model=List[dict])
async def get_reports(case_id: str, request: Request):
    """Get all reports for a case"""
    user = await get_current_user(request)
    
    reports = await db.reports.find(
        {"case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    ).sort("generated_at", -1).to_list(100)
    
    return reports

@api_router.get("/cases/{case_id}/reports/{report_id}", response_model=dict)
async def get_report(case_id: str, report_id: str, request: Request):
    """Get a specific report"""
    user = await get_current_user(request)
    
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return report

@api_router.delete("/cases/{case_id}/reports/{report_id}")
async def delete_report(case_id: str, report_id: str, request: Request):
    """Delete a report"""
    user = await get_current_user(request)
    
    result = await db.reports.delete_one({
        "report_id": report_id,
        "case_id": case_id,
        "user_id": user.user_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return {"message": "Report deleted"}

# ============ PDF EXPORT ============

@api_router.get("/cases/{case_id}/reports/{report_id}/export-pdf")
async def export_report_pdf(case_id: str, report_id: str, request: Request):
    """Export a report as PDF with Grounds of Merit and Legal References"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from io import BytesIO
    
    user = await get_current_user(request)
    
    # Get report
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get grounds of merit for this case
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Create PDF buffer
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=20*mm,
        leftMargin=20*mm,
        topMargin=20*mm,
        bottomMargin=20*mm
    )
    
    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ReportTitle',
        fontSize=24,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    ))
    styles.add(ParagraphStyle(
        name='ReportSubtitle',
        fontSize=14,
        spaceAfter=12,
        alignment=TA_CENTER,
        textColor=colors.grey
    ))
    styles.add(ParagraphStyle(
        name='SectionHeader',
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1e293b')
    ))
    styles.add(ParagraphStyle(
        name='ReportBodyText',
        fontSize=10,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
        leading=14
    ))
    styles.add(ParagraphStyle(
        name='LawSection',
        fontSize=9,
        spaceAfter=4,
        leftIndent=20,
        textColor=colors.HexColor('#1e40af')
    ))
    styles.add(ParagraphStyle(
        name='GroundTitle',
        fontSize=11,
        spaceBefore=10,
        spaceAfter=4,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#92400e')
    ))
    
    story = []
    
    # Header
    story.append(Paragraph("Criminal Appeal AI", styles['ReportSubtitle']))
    story.append(Paragraph("Prepared for: Deb King, Glenmore Park 2745", styles['ReportSubtitle']))
    story.append(Paragraph("One woman's fight for justice — seeking truth for Joshua Homann, failed by the system", styles['ReportSubtitle']))
    story.append(Paragraph("Criminal Appeal Case Management", styles['ReportSubtitle']))
    story.append(Spacer(1, 10*mm))
    
    # Report Title
    report_type_labels = {
        'quick_summary': 'Quick Case Summary',
        'full_detailed': 'Full Detailed Legal Analysis', 
        'extensive_log': 'Extensive Case Log & Analysis'
    }
    story.append(Paragraph(report_type_labels.get(report.get('report_type'), 'Legal Report'), styles['ReportTitle']))
    story.append(Spacer(1, 5*mm))
    
    # Case Info Table
    case_data_table = [
        ['Case Title:', case.get('title', 'N/A')],
        ['Defendant:', case.get('defendant_name', 'N/A')],
        ['Case Number:', case.get('case_number', 'N/A') or 'N/A'],
        ['Court:', case.get('court', 'N/A') or 'N/A'],
        ['Generated:', report.get('generated_at', 'N/A')[:10] if report.get('generated_at') else 'N/A']
    ]
    
    case_table = Table(case_data_table, colWidths=[40*mm, 120*mm])
    case_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#475569')),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(case_table)
    story.append(Spacer(1, 10*mm))
    
    # Grounds of Merit Section
    if grounds:
        story.append(Paragraph("GROUNDS OF MERIT", styles['SectionHeader']))
        story.append(Spacer(1, 5*mm))
        
        ground_type_labels = {
            'procedural_error': 'Procedural Error',
            'fresh_evidence': 'Fresh Evidence',
            'miscarriage_of_justice': 'Miscarriage of Justice',
            'sentencing_error': 'Sentencing Error',
            'judicial_error': 'Judicial Error',
            'ineffective_counsel': 'Ineffective Counsel',
            'prosecution_misconduct': 'Prosecution Misconduct',
            'jury_irregularity': 'Jury Irregularity',
            'constitutional_violation': 'Constitutional Violation',
            'other': 'Other'
        }
        
        for idx, ground in enumerate(grounds, 1):
            ground_type = ground_type_labels.get(ground.get('ground_type'), 'Other')
            strength = ground.get('strength', 'moderate').capitalize()
            
            # Ground header
            story.append(Paragraph(
                f"{idx}. {ground.get('title', 'Unnamed Ground')} [{ground_type}] - Strength: {strength}",
                styles['GroundTitle']
            ))
            
            # Description
            if ground.get('description'):
                story.append(Paragraph(ground.get('description'), styles['ReportBodyText']))
            
            # Legal References (Law Sections)
            if ground.get('law_sections'):
                story.append(Paragraph("<b>Relevant Law Sections:</b>", styles['ReportBodyText']))
                for section in ground.get('law_sections', []):
                    section_text = f"• s.{section.get('section', '')} {section.get('act', '')} ({section.get('jurisdiction', 'NSW')})"
                    story.append(Paragraph(section_text, styles['LawSection']))
            
            # Similar Cases
            if ground.get('similar_cases'):
                story.append(Paragraph("<b>Similar Cases:</b>", styles['ReportBodyText']))
                for case_ref in ground.get('similar_cases', []):
                    case_text = f"• {case_ref.get('case_name', '')} {case_ref.get('citation', '')}"
                    story.append(Paragraph(case_text, styles['LawSection']))
            
            # Supporting Evidence
            if ground.get('supporting_evidence'):
                story.append(Paragraph("<b>Supporting Evidence:</b>", styles['ReportBodyText']))
                for evidence in ground.get('supporting_evidence', []):
                    story.append(Paragraph(f"• {evidence}", styles['LawSection']))
            
            story.append(Spacer(1, 5*mm))
    
    # Legal Framework Reference
    story.append(Paragraph("LEGAL FRAMEWORK REFERENCE", styles['SectionHeader']))
    legal_refs = [
        "• Crimes Act 1900 (NSW) - Primary criminal law for NSW",
        "• Criminal Appeal Act 1912 (NSW) - Governs appeals in NSW",
        "• Criminal Code Act 1995 (Cth) - Federal criminal law",
        "• Evidence Act 1995 (NSW & Cth) - Evidence admissibility",
        "• Sentencing Act 1995 (NSW) - Sentencing guidelines"
    ]
    for ref in legal_refs:
        story.append(Paragraph(ref, styles['LawSection']))
    
    story.append(Spacer(1, 10*mm))
    
    # Main Analysis Content
    story.append(Paragraph("DETAILED ANALYSIS", styles['SectionHeader']))
    
    analysis_text = report.get('content', {}).get('analysis', 'No analysis available.')
    
    # Split analysis into paragraphs for better formatting
    paragraphs = analysis_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Clean up markdown-style formatting
            clean_para = para.replace('**', '').replace('##', '').strip()
            if clean_para:
                story.append(Paragraph(clean_para, styles['ReportBodyText']))
                story.append(Spacer(1, 3*mm))
    
    # Footer disclaimer
    story.append(Spacer(1, 15*mm))
    story.append(Paragraph(
        "This report was generated by Criminal Appeal AI for Deb King, Glenmore Park 2745. It should be reviewed by qualified legal counsel before being relied upon in legal proceedings.",
        ParagraphStyle(
            name='Disclaimer',
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
    ))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    # Create filename
    safe_title = "".join(c for c in case.get('title', 'Report')[:30] if c.isalnum() or c in ' -_').strip()
    filename = f"{safe_title}_{report.get('report_type', 'report')}.pdf"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

@api_router.get("/cases/{case_id}/reports/{report_id}/export-docx")
async def export_report_docx(case_id: str, report_id: str, request: Request):
    """Export a report as DOCX (Microsoft Word) with Grounds of Merit and Legal References"""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from io import BytesIO
    
    user = await get_current_user(request)
    
    # Get report
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get grounds of merit
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Create DOCX document
    doc = DocxDocument()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(146, 64, 14)
    
    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("CRIMINAL APPEAL AI")
    header_run.font.size = Pt(12)
    header_run.font.color.rgb = RGBColor(100, 116, 139)
    
    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_header.add_run("Prepared for: Deb King, Glenmore Park 2745")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    
    motto_para = doc.add_paragraph()
    motto_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    motto_run = motto_para.add_run("One woman's fight for justice — seeking truth for Joshua Homann, failed by the system")
    motto_run.font.size = Pt(9)
    motto_run.font.italic = True
    motto_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    
    # Report Title
    report_type_labels = {
        'quick_summary': 'Quick Case Summary',
        'full_detailed': 'Full Detailed Legal Analysis', 
        'extensive_log': 'Extensive Case Log & Analysis'
    }
    title = doc.add_heading(report_type_labels.get(report.get('report_type'), 'Legal Report'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Case Information Table
    case_table = doc.add_table(rows=5, cols=2)
    case_table.style = 'Table Grid'
    
    case_info = [
        ('Case Title:', case.get('title', 'N/A')),
        ('Defendant:', case.get('defendant_name', 'N/A')),
        ('Case Number:', case.get('case_number', 'N/A') or 'N/A'),
        ('Court:', case.get('court', 'N/A') or 'N/A'),
        ('Generated:', report.get('generated_at', 'N/A')[:10] if report.get('generated_at') else 'N/A')
    ]
    
    for i, (label, value) in enumerate(case_info):
        row = case_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Grounds of Merit Section
    if grounds:
        doc.add_heading('GROUNDS OF MERIT', level=1)
        
        ground_type_labels = {
            'procedural_error': 'Procedural Error',
            'fresh_evidence': 'Fresh Evidence',
            'miscarriage_of_justice': 'Miscarriage of Justice',
            'sentencing_error': 'Sentencing Error',
            'judicial_error': 'Judicial Error',
            'ineffective_counsel': 'Ineffective Counsel',
            'prosecution_misconduct': 'Prosecution Misconduct',
            'jury_irregularity': 'Jury Irregularity',
            'constitutional_violation': 'Constitutional Violation',
            'other': 'Other'
        }
        
        for idx, ground in enumerate(grounds, 1):
            ground_type = ground_type_labels.get(ground.get('ground_type'), 'Other')
            strength = ground.get('strength', 'moderate').capitalize()
            
            # Ground heading
            doc.add_heading(
                f"{idx}. {ground.get('title', 'Unnamed Ground')} [{ground_type}] - Strength: {strength}",
                level=2
            )
            
            # Description
            if ground.get('description'):
                doc.add_paragraph(ground.get('description'))
            
            # Legal References
            if ground.get('law_sections'):
                law_para = doc.add_paragraph()
                law_run = law_para.add_run('Relevant Law Sections:')
                law_run.font.bold = True
                law_run.font.color.rgb = RGBColor(30, 64, 175)
                
                for section in ground.get('law_sections', []):
                    section_text = f"s.{section.get('section', '')} {section.get('act', '')} ({section.get('jurisdiction', 'NSW')})"
                    bullet = doc.add_paragraph(section_text, style='List Bullet')
                    for run in bullet.runs:
                        run.font.color.rgb = RGBColor(30, 64, 175)
            
            # Similar Cases
            if ground.get('similar_cases'):
                cases_para = doc.add_paragraph()
                cases_run = cases_para.add_run('Similar Cases:')
                cases_run.font.bold = True
                cases_run.font.color.rgb = RGBColor(146, 64, 14)
                
                for case_ref in ground.get('similar_cases', []):
                    case_text = f"{case_ref.get('case_name', '')} {case_ref.get('citation', '')}"
                    doc.add_paragraph(case_text, style='List Bullet')
            
            # Supporting Evidence
            if ground.get('supporting_evidence'):
                evidence_para = doc.add_paragraph()
                evidence_run = evidence_para.add_run('Supporting Evidence:')
                evidence_run.font.bold = True
                evidence_run.font.color.rgb = RGBColor(5, 150, 105)
                
                for evidence in ground.get('supporting_evidence', []):
                    doc.add_paragraph(evidence, style='List Bullet')
            
            doc.add_paragraph()
    
    # Legal Framework Reference
    doc.add_heading('LEGAL FRAMEWORK REFERENCE', level=1)
    
    legal_refs = [
        "Crimes Act 1900 (NSW) - Primary criminal law for New South Wales",
        "Criminal Appeal Act 1912 (NSW) - Governs criminal appeals in NSW",
        "Criminal Code Act 1995 (Cth) - Federal criminal law",
        "Evidence Act 1995 (NSW & Cth) - Rules on evidence admissibility",
        "Sentencing Act 1995 (NSW) - Sentencing guidelines and procedures"
    ]
    
    for ref in legal_refs:
        doc.add_paragraph(ref, style='List Bullet')
    
    doc.add_paragraph()
    
    # Detailed Analysis
    doc.add_heading('DETAILED ANALYSIS', level=1)
    
    analysis_text = report.get('content', {}).get('analysis', 'No analysis available.')
    
    # Split analysis into paragraphs
    paragraphs = analysis_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            clean_para = para.replace('**', '').replace('##', '').strip()
            if clean_para:
                doc.add_paragraph(clean_para)
    
    # Footer disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer.add_run(
        "This report was generated by Justitia AI and should be reviewed by qualified legal counsel before being relied upon in legal proceedings."
    )
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.italic = True
    disclaimer_run.font.color.rgb = RGBColor(100, 116, 139)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create filename
    safe_title = "".join(c for c in case.get('title', 'Report')[:30] if c.isalnum() or c in ' -_').strip()
    filename = f"{safe_title}_{report.get('report_type', 'report')}.docx"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

# ============ HEALTH CHECK ============

@api_router.get("/")
async def root():
    return {"message": "Criminal Appeal AI API", "status": "operational"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
